# ============================================================
# views.py — ConformExpert
# ============================================================

from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.template.loader import render_to_string
from django.core.mail import send_mail
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from django.core.paginator import Paginator
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status

from .models import Document, DocumentFile, Analysis, Devis, FactureEnergie, Message, Avis
from .forms import DocumentForm, ContactForm
from .serializers import DocumentSerializer, AnalysisSerializer

import PyPDF2
import re
import threading
import base64
import json
import csv
import calendar
from datetime import timedelta, date
from collections import defaultdict


# ──────────────────────────────────────────────────────────────
# CONSTANTES
# ──────────────────────────────────────────────────────────────

SITE_URL = "https://conformexpert.cc"


# ──────────────────────────────────────────────────────────────
# EMAILS — helpers internes
# ──────────────────────────────────────────────────────────────

def _send_html_async(sujet, template_name, context, destinataire):
    """Envoie un email HTML via SendGrid dans un thread séparé."""
    if not destinataire:
        return

    def _send():
        try:
            import sendgrid
            from sendgrid.helpers.mail import Mail

            html = render_to_string(f'main/emails/{template_name}', context)
            sg = sendgrid.SendGridAPIClient(api_key=settings.SENDGRID_API_KEY)
            message = Mail(
                from_email=settings.DEFAULT_FROM_EMAIL,
                to_emails=destinataire,
                subject=sujet,
                html_content=html,
            )
            response = sg.send(message)
            print(f"MAIL SENDGRID OK → {destinataire} (status {response.status_code})")
        except Exception as e:
            print(f"ERREUR MAIL : {e}")

    t = threading.Thread(target=_send, daemon=True)
    t.start()


def send_mail_reception(document):
    if not document.client_email:
        return
    _send_html_async(
        f"[ConformExpert] Dossier bien reçu — {document.name}",
        "email_reception.html",
        {
            'doc_id': f"{document.id:04d}",
            'doc_name': document.name,
            'client_name': document.client_name or '',
            'date_depot': document.upload_date.strftime('%d/%m/%Y à %H:%M'),
        },
        document.client_email,
    )


def send_mail_validation_devis(document, devis=None):
    if not document.client_email:
        return

    montant_ht = float(devis.montant) if devis and devis.montant else 0
    tva = round(montant_ht * 0.20, 2)

    _send_html_async(
        f"[ConformExpert] Votre devis — {document.name}",
        "email_devis.html",
        {
            'doc_id': f"{document.id:04d}",
            'doc_name': document.name,
            'client_name': document.client_name or '',
            'accepter_url': f"{SITE_URL}/devis/accepter/{devis.id}/",
            'refuser_url': f"{SITE_URL}/devis/refuser/{devis.id}/",
            'montant_ht': f"{montant_ht:.2f}",
            'tva': f"{tva:.2f}",
            'montant_ttc': f"{montant_ht + tva:.2f}",
            'norme': devis.norme if devis else 'RT2012 / RE2020',
            'notes': devis.notes if devis else '',
        },
        document.client_email,
    )


def send_mail_analyse_commence(document):
    if not document.client_email:
        return
    _send_html_async(
        "[ConformExpert] L'analyse de votre dossier a démarré",
        "email_analyse_commence.html",
        {
            'doc_id': f"{document.id:04d}",
            'doc_name': document.name,
            'client_name': document.client_name or '',
            'tracking_url': f"{SITE_URL}/suivi/{document.tracking_token}/",
        },
        document.client_email,
    )


def send_mail_analyse_terminee(document):
    if not document.client_email:
        return
    _send_html_async(
        f"[ConformExpert] Votre rapport est disponible — {document.name}",
        "email_analyse_terminee.html",
        {
            'doc_id': f"{document.id:04d}",
            'doc_name': document.name,
            'client_name': document.client_name or '',
            'tracking_url': f"{SITE_URL}/suivi/{document.tracking_token}/",
            'rapport_items': [
                "Analyse des critères thermiques et énergétiques",
                "Conclusion de conformité détaillée",
                "Recommandations éventuelles",
                "Rapport PDF téléchargeable",
            ],
        },
        document.client_email,
    )


# ──────────────────────────────────────────────────────────────
# AVIS CLIENTS
# ──────────────────────────────────────────────────────────────

def send_mail_avis(document):
    """
    Crée (ou récupère) l'objet Avis du dossier et envoie
    l'email d'invitation une seule fois.
    Appelable manuellement (bouton admin) ou automatiquement
    à la fin de send_mail_analyse_terminee().
    """
    if not document.client_email:
        return

    avis, _ = Avis.objects.get_or_create(document=document)

    if avis.certifie or avis.email_envoye:
        return  # déjà noté ou déjà envoyé → ne rien faire

    _send_html_async(
        f"[ConformExpert] Donnez votre avis sur votre analyse — {document.name}",
        "email_avis.html",
        {
            'doc_id':      f"{document.id:04d}",
            'doc_name':    document.name,
            'client_name': document.client_name or '',
            'avis_url':    f"{SITE_URL}/avis/{avis.token}/",
        },
        document.client_email,
    )

    avis.email_envoye = True
    avis.save(update_fields=['email_envoye'])


def noter_service(request, token):
    """
    Page publique /avis/<token>/
    GET  → formulaire (avec note pré-sélectionnée si ?note=X)
    POST → enregistre et marque certifié
    """
    avis   = get_object_or_404(Avis, token=token)
    doc_id = f"{avis.document.id:04d}"

    # Déjà noté → confirmation lecture seule
    if avis.certifie:
        return render(request, 'main/avis_noter.html', {
            'avis':   avis,
            'doc_id': doc_id,
        })

    erreur      = None
    note_pre    = request.GET.get('note') or ''
    commentaire = ''

    if request.method == 'POST':
        note_pre    = request.POST.get('note', '')
        commentaire = request.POST.get('commentaire', '').strip()

        try:
            note = int(note_pre)
        except (ValueError, TypeError):
            note = 0

        if note not in range(1, 6):
            erreur = "Veuillez sélectionner une note entre 1 et 5 étoiles."
        else:
            avis.note        = note
            avis.commentaire = commentaire[:1000]
            avis.certifie    = True
            avis.soumis_le   = timezone.now()
            avis.save(update_fields=['note', 'commentaire', 'certifie', 'soumis_le'])

            return render(request, 'main/avis_noter.html', {
                'avis':   avis,
                'doc_id': doc_id,
            })

    return render(request, 'main/avis_noter.html', {
        'avis':            avis,
        'doc_id':          doc_id,
        'erreur':          erreur,
        'note_pre':        note_pre,
        'commentaire_pre': commentaire,
    })


@login_required
def envoyer_invitation_avis(request, doc_id):
    """
    Appelée depuis le bouton ⭐ dans l'onglet Communication
    de edit_document.html.
    Force le renvoi (relance) en remettant email_envoye à False.
    """
    document = get_object_or_404(Document, id=doc_id)

    if not document.client_email:
        messages.error(request, "Aucun email client renseigné pour ce dossier.")
        return redirect('edit_document', doc_id=doc_id)

    avis, _ = Avis.objects.get_or_create(document=document)

    if avis.certifie:
        messages.warning(request, "Ce client a déjà laissé son avis — impossible de renvoyer.")
        return redirect('edit_document', doc_id=doc_id)

    # Forcer le renvoi même si email_envoye=True (relance)
    avis.email_envoye = False
    avis.save(update_fields=['email_envoye'])

    send_mail_avis(document)
    messages.success(request, f"Invitation à noter envoyée à {document.client_email}.")
    return redirect('edit_document', doc_id=doc_id)


def avis_publics(request):
    """
    Rendu du widget avis pour la landing page (vue standalone).
    """
    from django.db.models import Avg, Count

    stats = Avis.objects.filter(certifie=True, note__isnull=False).aggregate(
        moyenne=Avg('note'),
        total=Count('id'),
    )
    avis_list = (
        Avis.objects
        .filter(certifie=True)
        .select_related('document')
        .order_by('-soumis_le')[:20]
    )

    return render(request, 'main/avis_widget.html', {
        'avis_list': avis_list,
        'moyenne':   round(stats['moyenne'] or 0, 1),
        'total':     stats['total'],
    })


# ──────────────────────────────────────────────────────────────
# DEVIS — acceptation / refus publics
# ──────────────────────────────────────────────────────────────

def accepter_devis(request, devis_id):
    devis = get_object_or_404(Devis, id=devis_id)

    if devis.statut != "accepte":
        devis.statut = "accepte"
        devis.save()

        if devis.document:
            devis.document.status = "en_cours"
            devis.document.save()

        _send_html_async(
            "✅ Devis accepté — ConformExpert",
            "email_notification_admin.html",
            {
                "sujet":        "Devis accepté par le client",
                "client":       devis.client_nom,
                "projet":       devis.projet_nom,
                "montant":      devis.montant,
                "doc_id":       f"{devis.id:04d}",
                "type_analyse": devis.document.type_analyse if devis.document else None,
                "admin_url":    f"{SITE_URL}/dossier/{devis.document.id}/editer/" if devis.document else None,
            },
            "contact@conformexpert.cc",
        )

    return render(request, "main/devis_accepte.html", {"devis": devis})


def refuser_devis(request, devis_id):
    devis = get_object_or_404(Devis, id=devis_id)

    if devis.statut != "refuse":
        devis.statut = "refuse"
        devis.save()

        _send_html_async(
            "❌ Devis refusé — ConformExpert",
            "email_notification_admin.html",
            {
                "sujet":        "Devis refusé par le client",
                "client":       devis.client_nom,
                "projet":       devis.projet_nom,
                "montant":      devis.montant,
                "doc_id":       f"{devis.id:04d}",
                "type_analyse": devis.document.type_analyse if devis.document else None,
            },
            "contact@conformexpert.cc",
        )

    return render(request, "main/devis_refuse.html", {"devis": devis})


# ──────────────────────────────────────────────────────────────
# AUTH
# ──────────────────────────────────────────────────────────────

def admin_login(request):
    if request.user.is_authenticated:
        return redirect('home')
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None and (user.is_staff or user.is_superuser):
            login(request, user)
            return redirect(request.GET.get('next', 'home'))
        messages.error(request, 'Identifiants incorrects ou accès non autorisé.')
    return render(request, 'main/login.html')


def admin_logout(request):
    logout(request)
    return redirect('landing')


# ──────────────────────────────────────────────────────────────
# RÉFÉRENTIELS RÉGLEMENTAIRES
# ──────────────────────────────────────────────────────────────

def fetch_re2020_requirements():
    return {
        'energy_efficiency': 80.0,
        'thermal_comfort': 85.0,
        'carbon_emissions': 75.0,
        'water_management': 70.0,
        'indoor_air_quality': 75.0,
    }


def fetch_rt2012_requirements():
    return {
        'bbio': 50.0,
        'cep': 50.0,
        'tic': 27.0,
        'airtightness': 0.6,
        'enr': 1.0,
    }


# ──────────────────────────────────────────────────────────────
# EXTRACTION / ANALYSE PDF
# ──────────────────────────────────────────────────────────────

def extract_text_from_pdf(upload_path):
    """Extrait le texte brut d'un fichier PDF via PyPDF2."""
    text = ""
    try:
        with open(upload_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
    except Exception as e:
        print(f"Erreur lecture PDF: {e}")
    return text


# ──────────────────────────────────────────────────────────────
# PARSER INTELLIGENT — détection + extraction + validation
# ──────────────────────────────────────────────────────────────

_PROMPT_DETECTION = """Tu es un expert en réglementation thermique française.
Analyse ce texte extrait d'un document thermique et réponds UNIQUEMENT en JSON valide, sans markdown.

Détermine :
1. Le type de document
2. Le logiciel utilisé pour le produire
3. La norme applicable
4. Toutes les valeurs thermiques présentes
5. Les métadonnées du bâtiment
6. Les alertes de cohérence éventuelles

JSON attendu (inclus uniquement les clés dont tu as trouvé la valeur) :
{
  "type_rapport": "climawin_rt2012" | "climawin_re2020" | "pleiades_rt2012" | "pleiades_re2020" | "dpe" | "attestation_rt2012" | "attestation_re2020" | "etude_thermique" | "autre",
  "logiciel_detecte": "ex: Climawin v4.2 / Pléiades+Comfie 7.1 / non détecté",
  "version_norme_detectee": "ex: RT2012 - Arrêté 26/10/2010",
  "norme_suggeree": "RT2012" | "RE2020" | "PEB" | "MINERGIE" | "SIA380" | "CNEB2015" | "CNEB2020" | "LENOZ",

  "valeurs": {
    "rt2012_bbio": null,
    "rt2012_bbio_max": null,
    "rt2012_cep": null,
    "rt2012_cep_max": null,
    "rt2012_tic": null,
    "rt2012_tic_max": null,
    "rt2012_airtightness": null,
    "rt2012_enr": null,
    "re2020_energy_efficiency": null,
    "re2020_energy_efficiency_max": null,
    "re2020_thermal_comfort": null,
    "re2020_thermal_comfort_max": null,
    "re2020_carbon_emissions": null,
    "re2020_carbon_emissions_max": null,
    "dpe_classe_energie": null,
    "dpe_classe_ges": null,
    "dpe_conso_ep": null,
    "dpe_emission_ges": null,
    "dpe_surface_ref": null,
    "dpe_date_visite": null,
    "dpe_diagnostiqueur": null
  },

  "batiment": {
    "nom_projet": null,
    "adresse": null,
    "surface_totale": null,
    "annee_construction": null,
    "type_batiment": null,
    "zone_climatique": null,
    "maitre_ouvrage": null,
    "bureau_etudes": null,
    "date_etude": null
  },

  "conformite_declaree": "conforme" | "non_conforme" | "non_precise",

  "alertes": [
    "Description d'une anomalie ou incohérence détectée dans les valeurs"
  ],

  "resume_extraction": "Phrase courte décrivant ce qui a été trouvé dans le document"
}

N'inclus dans "valeurs" que les clés dont la valeur est non nulle.
Dans "alertes", signale : valeurs manquantes importantes, incohérences entre valeurs déclarées et seuils, valeurs hors plage réaliste.

Texte du document :
"""

_PROMPT_DETECTION_SUFFIX = "\n\nRéponds UNIQUEMENT avec le JSON, sans texte avant ni après, sans balises markdown."


def analyser_rapport_thermique(texte, pdf_b64=None):
    """
    Analyse intelligente d'un rapport thermique via Claude.
    Détecte le type (Climawin RT2012, Pleiades RE2020, DPE…),
    extrait toutes les valeurs et génère des alertes de cohérence.
    Retourne un dict avec type_rapport, valeurs, alertes, métadonnées.
    """
    import os
    import urllib.request
    import urllib.error

    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
    if not ANTHROPIC_API_KEY:
        return _fallback_regex(texte)

    try:
        # Construire le message — avec PDF natif si disponible
        user_content = []
        headers_extra = {}

        if pdf_b64:
            user_content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64},
            })
            headers_extra = {"anthropic-beta": "pdfs-2024-09-25"}
            user_content.append({
                "type": "text",
                "text": _PROMPT_DETECTION + "(document PDF joint)" + _PROMPT_DETECTION_SUFFIX,
            })
        else:
            user_content.append({
                "type": "text",
                "text": _PROMPT_DETECTION + texte[:12000] + _PROMPT_DETECTION_SUFFIX,
            })

        payload = json.dumps({
            "model": "claude-sonnet-4-5",
            "max_tokens": 4000,
            "messages": [{"role": "user", "content": user_content}],
        }).encode('utf-8')

        headers = {
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
        }
        headers.update(headers_extra)

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload, headers=headers, method="POST",
        )
        with urllib.request.urlopen(req, timeout=300) as resp:
            result = json.loads(resp.read().decode('utf-8'))

            raw = result['content'][0]['text'].strip()

            # enlever markdown éventuel
            raw = raw.replace('```json', '').replace('```', '').strip()

            # extraire uniquement le JSON
            match = re.search(r'\{[\s\S]*\}', raw)
            if match:
                raw = match.group(0)
            else:
                print("❌ Aucun JSON détecté")
                return _fallback_regex(texte)

            try:
                data = json.loads(raw)

                print(f"PARSER OK — type={data.get('type_rapport')} norme={data.get('norme_suggeree')}")
                return data

            except Exception as json_err:
                print("❌ JSON parsing error :", json_err)
                print("RAW PREVIEW:", raw[:1000])
                return _fallback_regex(texte)

    except Exception as e:
        print("❌ Erreur API Claude :", e)
        try:
            print("RAW PREVIEW:", raw[:1000])
        except:
            pass
        return _fallback_regex(texte)


def _fallback_regex(texte):
    """Fallback regex minimal si l'API Claude est indisponible."""
    data = {'type_rapport': 'inconnu', 'valeurs': {}, 'alertes': [], 'logiciel_detecte': ''}

    # Détection logiciel
    t_low = texte.lower()
    if 'climawin' in t_low:
        data['logiciel_detecte'] = 'Climawin'
    elif 'pleiades' in t_low or 'pléiades' in t_low:
        data['logiciel_detecte'] = 'Pléiades'
    elif 'dpe' in t_low or 'diagnostic de performance' in t_low:
        data['logiciel_detecte'] = 'DPE'

    # Détection norme
    if 're2020' in t_low or 're 2020' in t_low:
        data['norme_suggeree'] = 'RE2020'
        data['type_rapport'] = 'climawin_re2020' if 'climawin' in t_low else 'pleiades_re2020' if 'pleiades' in t_low else 'etude_thermique'
    elif 'rt2012' in t_low or 'rt 2012' in t_low:
        data['norme_suggeree'] = 'RT2012'
        data['type_rapport'] = 'climawin_rt2012' if 'climawin' in t_low else 'pleiades_rt2012' if 'pleiades' in t_low else 'etude_thermique'
    elif 'dpe' in t_low:
        data['norme_suggeree'] = 'RE2020'
        data['type_rapport'] = 'dpe'

    # Extraction regex basique
    valeurs = {}
    for pattern, key in [
        (r'Bbio\s*[=:]\s*([\d.,]+)',                 'rt2012_bbio'),
        (r'Cep\s*[=:]\s*([\d.,]+)',                  'rt2012_cep'),
        (r'Tic\s*[=:]\s*([\d.,]+)',                  'rt2012_tic'),
        (r'[Ee]tanch[eé]it[eé]\s*[=:]\s*([\d.,]+)', 'rt2012_airtightness'),
        (r'ENR\s*[=:]\s*([\d.,]+)',                  'rt2012_enr'),
        (r'Cep,?nr\s*[=:]\s*([\d.,]+)',              're2020_energy_efficiency'),
        (r'DH\s*[=:]\s*([\d.,]+)',                   're2020_thermal_comfort'),
        (r'Ic.{0,10}[ée]nergie\s*[=:]\s*([\d.,]+)', 're2020_carbon_emissions'),
        (r'classe\s*[=:]\s*([A-G])',                 'dpe_classe_energie'),
    ]:
        m = re.search(pattern, texte, re.IGNORECASE)
        if m:
            val = m.group(1).replace(',', '.')
            try:
                valeurs[key] = float(val) if key != 'dpe_classe_energie' else val
            except ValueError:
                pass
    data['valeurs'] = valeurs
    return data


def parse_pdf_text(text, norme=None):
    """
    Compatibilité ascendante — appelle le nouveau parser et retourne
    le format dict attendu par analyze_document.
    """
    result = analyser_rapport_thermique(text)
    return result.get('valeurs', {})


def analyze_document(document, data, resultat_complet=None):
    """
    Hydrate les champs thermiques + métadonnées d'un document.
    data = dict valeurs thermiques
    resultat_complet = dict complet retourné par analyser_rapport_thermique
    """
    # ── Valeurs thermiques ─────────────────────────────────────
    document.rt2012_bbio         = data.get('rt2012_bbio')
    document.rt2012_cep          = data.get('rt2012_cep')
    document.rt2012_tic          = data.get('rt2012_tic')
    document.rt2012_airtightness = data.get('rt2012_airtightness')
    document.rt2012_enr          = data.get('rt2012_enr')
    document.re2020_energy_efficiency = data.get('re2020_energy_efficiency')
    document.re2020_thermal_comfort   = data.get('re2020_thermal_comfort')
    document.re2020_carbon_emissions  = data.get('re2020_carbon_emissions')
    document.peb_espec      = data.get('peb_espec')
    document.peb_ew         = data.get('peb_ew')
    document.peb_u_mur      = data.get('peb_u_mur')
    document.peb_u_toit     = data.get('peb_u_toit')
    document.peb_u_plancher = data.get('peb_u_plancher')
    document.minergie_qh    = data.get('minergie_qh')
    document.minergie_qtot  = data.get('minergie_qtot')
    document.minergie_n50   = data.get('minergie_n50')
    document.sia380_qh      = data.get('sia380_qh') or data.get('minergie_qh')
    document.cneb_ei           = data.get('cneb_ei')
    document.cneb_u_mur        = data.get('cneb_u_mur')
    document.cneb_u_toit       = data.get('cneb_u_toit')
    document.cneb_u_fenetre    = data.get('cneb_u_fenetre')
    document.cneb_infiltration = data.get('cneb_infiltration')
    document.lenoz_ep     = data.get('lenoz_ep')
    document.lenoz_ew     = data.get('lenoz_ew')
    document.lenoz_u_mur  = data.get('lenoz_u_mur')
    document.lenoz_u_toit = data.get('lenoz_u_toit')

    # ── Champs DPE ─────────────────────────────────────────────
    if data.get('dpe_classe_energie'):
        document.dpe_classe_energie = str(data['dpe_classe_energie']).upper()
    if data.get('dpe_classe_ges'):
        document.dpe_classe_ges = str(data['dpe_classe_ges']).upper()
    if data.get('dpe_conso_ep'):
        document.dpe_conso_ep = data['dpe_conso_ep']
    if data.get('dpe_emission_ges'):
        document.dpe_emission_ges = data['dpe_emission_ges']
    if data.get('dpe_surface_ref'):
        document.dpe_surface_ref = data['dpe_surface_ref']
    if data.get('dpe_date_visite'):
        document.dpe_date_visite = str(data['dpe_date_visite'])
    if data.get('dpe_diagnostiqueur'):
        document.dpe_diagnostiqueur = str(data['dpe_diagnostiqueur'])

    # ── Métadonnées issues du parser complet ───────────────────
    if resultat_complet:
        document.type_rapport       = resultat_complet.get('type_rapport', 'inconnu')
        document.logiciel_detecte   = resultat_complet.get('logiciel_detecte', '')
        document.version_norme_detectee = resultat_complet.get('version_norme_detectee', '')
        document.extraction_ok      = bool(resultat_complet.get('valeurs'))
        document.extraction_json    = resultat_complet
        document.extraction_alertes = resultat_complet.get('alertes', [])

        # Auto-mise à jour de la norme si détectée
        norme_suggeree = resultat_complet.get('norme_suggeree')
        if norme_suggeree and document.norme == 'RE2020':
            document.norme = norme_suggeree

        # Auto-remplissage métadonnées bâtiment si vides
        bat = resultat_complet.get('batiment', {}) or {}
        if bat.get('surface_totale') and not document.surface_totale:
            try:
                document.surface_totale = float(bat['surface_totale'])
            except (ValueError, TypeError):
                pass
        if bat.get('annee_construction') and not document.annee_construction:
            try:
                document.annee_construction = int(bat['annee_construction'])
            except (ValueError, TypeError):
                pass

    document.save()




# ──────────────────────────────────────────────────────────────
# VUES PUBLIQUES
# ──────────────────────────────────────────────────────────────


def maintenance(request):
    """Page de maintenance — retourne 503."""
    from django.template.response import TemplateResponse
    return TemplateResponse(request, 'main/maintenance.html', status=503)

def landing(request):
    """Page d'accueil publique."""
    from django.db.models import Avg, Count
    stats = Avis.objects.filter(certifie=True, note__isnull=False).aggregate(
        moyenne=Avg('note'),
        total=Count('id'),
    )
    avis_list = (
        Avis.objects
        .filter(certifie=True)
        .select_related('document')
        .order_by('-soumis_le')[:9]
    )
    return render(request, 'main/landing.html', {
        'avis_list': avis_list,
        'moyenne':   round(stats['moyenne'] or 0, 1),
        'total':     stats['total'],
    })


def contact(request):
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            try:
                send_mail(
                    subject=f"[ConformExpert] Nouveau contact : {form.cleaned_data['name']}",
                    message=(
                        f"Nom : {form.cleaned_data['name']}\n"
                        f"Email : {form.cleaned_data['email']}\n"
                        f"Téléphone : {form.cleaned_data.get('phone', 'N/A')}\n"
                        f"Profil : {form.cleaned_data.get('profile', 'N/A')}\n\n"
                        f"Message :\n{form.cleaned_data['message']}"
                    ),
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[settings.CONTACT_EMAIL],
                    fail_silently=True,
                )
            except Exception:
                pass
            messages.success(request, 'Message envoyé. Nous vous répondons sous 48h.')
            return redirect('contact')
    else:
        form = ContactForm()
    return render(request, 'main/contact.html', {'form': form})


def mentions_legales(request):
    return render(request, 'main/mentions_legales.html')


def faq(request):
    faq_items = [
        {
            "question": "ConformExpert peut-il contredire le rapport de mon bureau d'études ?",
            "answer": (
                "Oui — c'est précisément notre rôle. ConformExpert est un tiers indépendant. "
                "Si nous détectons une valeur incohérente, un seuil non respecté ou une hypothèse discutable, "
                "nous le signalons explicitement dans notre rapport de validation, avec le niveau de criticité correspondant. "
                "Nous ne sommes liés à aucun bureau d'études."
            ),
        },
        {
            "question": "Quels fichiers dois-je fournir ?",
            "answer": (
                "Le rapport thermique ou énergétique de votre bureau d'études en PDF ou XML (Climawin, Pléiades, DPE, PEB, CNEB…). "
                "Si disponibles, les factures énergie permettent le croisement avec les consommations modélisées. "
                "Pour un bilan carbone : DPE, attestation RE2020, factures énergie. "
                "Plus le dossier est complet, plus la validation est précise."
            ),
        },
        {
            "question": "Le rapport de validation a-t-il une valeur légale ?",
            "answer": (
                "Notre rapport de validation est un avis d'expert indépendant. Il peut être utilisé en support d'une négociation, "
                "d'un recours amiable ou d'une procédure judiciaire, comme élément de preuve attestant d'une incohérence "
                "ou d'une non-conformité identifiée. Il ne se substitue pas à un acte réglementaire officiel."
            ),
        },
        {
            "question": "Puis-je déposer un dossier sans avoir de rapport complet ?",
            "answer": (
                "Oui — nous pouvons travailler sur un rapport partiel ou en cours de finalisation. "
                "Contactez-nous avant le dépôt pour évaluer ensemble ce qui est nécessaire à une validation complète ou partielle."
            ),
        },
        {
            "question": "Intervenez-vous sur des projets hors de France ?",
            "answer": (
                "Oui. ConformExpert traite les dossiers PEB (Belgique), CNEB (Canada) et LENOZ (Luxembourg) "
                "en plus des normes françaises RT2012 et RE2020. Pour d'autres normes, contactez-nous pour évaluer la faisabilité."
            ),
        },
        {
            "question": "Quelle est la différence entre RT2012 et RE2020 ?",
            "answer": (
                "La RT2012 encadre la consommation énergétique via Bbio, Cep et Tic. "
                "La RE2020, en vigueur depuis janvier 2022, va plus loin : elle intègre le bilan carbone "
                "sur le cycle de vie du bâtiment (Ic énergie, Ic construction) et renforce les exigences de confort d'été."
            ),
        },
        {
            "question": "Quel est le délai de livraison ?",
            "answer": (
                "Nous garantissons la livraison du rapport sous 10 jours ouvrés après validation du devis et réception du paiement. "
                "Ce délai est affiché sur votre lien de suivi dès le démarrage de l'analyse."
            ),
        },
        {
            "question": "Comment fonctionne le lien de suivi ?",
            "answer": (
                "Après dépôt, vous recevez un lien unique personnalisé. Il vous permet de suivre l'avancement en temps réel "
                "et de télécharger le rapport dès sa livraison. Aucune création de compte n'est nécessaire."
            ),
        },
        {
            "question": "Mon analyse est-elle vraiment indépendante ?",
            "answer": (
                "Oui. Notre analyse est réalisée sans lien avec le bureau d'études ou le maître d'ouvrage. "
                "Cette indépendance garantit une lecture objective et non biaisée de vos documents."
            ),
        },
        {
            "question": "Qu'est-ce qu'un bilan carbone immobilier ?",
            "answer": (
                "Le bilan carbone immobilier évalue les émissions de CO2 liées à l'usage du bâtiment "
                "(chauffage, électricité, eau chaude) et, si applicable, à sa construction (matériaux). "
                "Nous analysons les indicateurs Ic énergie et Ic construction (RE2020), la classe DPE, "
                "les émissions GES, et identifions les leviers de réduction."
            ),
        },
    ]
    return render(request, 'main/faq.html', {'faq_items': faq_items})


# ──────────────────────────────────────────────────────────────
# IMPORT & TRACKING
# ──────────────────────────────────────────────────────────────

def import_document(request):
    if request.method == "POST":
        data = request.POST.copy()
        type_analyse = data.get("type_analyse")

        # Pour bilan carbone, les champs énergie ne sont pas requis
        if type_analyse == "carbone":
            data["climate_zone"] = ""
            data["norme"] = ""

        form = DocumentForm(data, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.type_analyse = type_analyse
            document.save()

            # Multi-upload
            for f in request.FILES.getlist("uploads"):
                DocumentFile.objects.create(
                    document=document,
                    fichier=f,
                    nom=f.name,
                    taille=f.size,
                )

            # Extraction du texte de tous les fichiers
            texte_complet = ""
            for doc_file in document.fichiers.all():
                try:
                    texte_complet += extract_text_from_pdf(doc_file.fichier.path) + "\n\n"
                except Exception:
                    pass

            # Fallback sur l'ancien champ upload
            if not texte_complet and document.upload:
                try:
                    texte_complet = extract_text_from_pdf(document.upload.path)
                except Exception:
                    texte_complet = ""

            # Pour bilan carbone, on ne parse pas comme un rapport thermique
            if type_analyse == 'carbone':
                document.extraction_alertes = []
                document.save(update_fields=['extraction_alertes'])
            else:
                resultat_complet = analyser_rapport_thermique(texte_complet)
                valeurs = resultat_complet.get('valeurs', {})
                analyze_document(document, valeurs, resultat_complet)
            send_mail_reception(document)

            messages.success(request, "Dossier reçu. Votre lien de suivi a été créé.")
            return redirect('tracking', token=document.tracking_token)
        else:
            messages.error(request, "Veuillez corriger les erreurs ci-dessous.")
    else:
        form = DocumentForm()

    return render(request, "main/import.html", {
        "form": form,
        "doc_items": [
            "Notice ou étude thermique (RT2012, RE2020, PEB, Minergie, CNEB, LENOZ…)",
            "Plans architecturaux (PDF)",
            "Tout document technique lié à la conformité énergétique du bâtiment",
            "DPE ou équivalent si disponible",
        ],
        "steps": [
            "Accusé de réception sous 24h",
            "Confirmation de complétude du dossier",
            "Analyse documentaire complète",
            "Livraison du rapport IA + lien de suivi",
        ],
    })


def get_tracking_steps(document):
    """Retourne la liste des étapes de suivi avec leur état (done / active / pending)."""
    if getattr(document, 'type_analyse', 'energie') == 'carbone':
        steps_def = [
            ("Dossier reçu et validé",           'recu'),
            ("Extraction des données carbone",   'en_cours'),
            ("Analyse des émissions & DPE",      'en_cours'),
            ("Rédaction du bilan carbone",       'en_cours'),
            ("Livraison du rapport PDF",         'termine'),
        ]
    else:
        steps_def = [
            ("Dossier reçu et validé",              'recu'),
            ("Analyse de l'enveloppe thermique",     'en_cours'),
            ("Vérification systèmes & attestations", 'en_cours'),
            ("Rédaction du rapport",                 'en_cours'),
            ("Livraison du rapport PDF",             'termine'),
        ]
    order = ['recu', 'en_cours', 'termine']
    current_idx = order.index(document.status)
    result = []
    for label, needed_status in steps_def:
        needed_idx = order.index(needed_status)
        if current_idx > needed_idx:
            state = 'done'
        elif current_idx == needed_idx:
            state = 'active'
        else:
            state = 'pending'
        result.append((label, state))
    return result


def tracking(request, token):
    document = get_object_or_404(Document, tracking_token=token)
    step_list = get_tracking_steps(document)
    progress_pct = {'recu': 15, 'en_cours': 60, 'termine': 100}.get(document.status, 15)

    devis_accepte = False
    if request.GET.get('accepter_devis') == '1' and document.status == 'recu':
        document.status = 'en_cours'
        document.save()
        send_mail_analyse_commence(document)
        step_list = get_tracking_steps(document)
        progress_pct = 60
        devis_accepte = True

    # Marquer les messages admin comme lus par le client
    document.messages.filter(auteur='admin', lu_client=False).update(lu_client=True)

    # Extraire le score IA depuis le rapport sauvegardé
    rapport_ia_score = None
    if document.rapport_ia_json:
        try:
            rapport_ia_data = json.loads(document.rapport_ia_json)
            rapport_ia_score = rapport_ia_data.get('score_global')
        except Exception:
            pass
    # Offset SVG pour le cercle de score (circumference ≈ 201 pour r=32)
    score_ring_offset = round(201 * (1 - (rapport_ia_score or 0) / 100), 1) if rapport_ia_score is not None else 201

    return render(request, 'main/tracking.html', {
        'document': document,
        'step_list': step_list,
        'progress_pct': progress_pct,
        'devis_accepte': devis_accepte,
        'rapport_ia_score': rapport_ia_score,
        'rapport_ia_score_offset': score_ring_offset,
    })


# ──────────────────────────────────────────────────────────────
# VUES ADMIN — tableau de bord
# ──────────────────────────────────────────────────────────────

@login_required(login_url='/login/')
def home(request):
    if not request.user.is_staff:
        return redirect('landing')

    documents = Document.objects.filter(is_active=True).order_by('-upload_date')

    # Sous-ensembles par type
    energie_docs = documents.filter(type_analyse="energie")
    carbone_docs = documents.filter(type_analyse="carbone")
    

    total_projects  = documents.count()
    compliant_count = sum(1 for doc in documents if doc.is_conform is True)
    compliance_rate = round(compliant_count / total_projects * 100, 1) if total_projects else 0
    pending_count   = documents.filter(status='recu').count()

    five_days_ago = timezone.now() - timedelta(days=5)
    old_pending   = documents.filter(status='recu', upload_date__lt=five_days_ago).count()

    try:
        recent_devis    = list(Devis.objects.all()[:5])
        devis_en_attente = Devis.objects.filter(statut='en_attente').count()
    except Exception:
        recent_devis     = []
        devis_en_attente = 0

    # ── Stats mensuelles (6 derniers mois) ──────────────────
    monthly_data = []
    today = timezone.now().date()
    for i in range(5, -1, -1):
        # Arithmétique correcte par mois (pas d'approximation en jours)
        month = (today.month - i - 1) % 12 + 1
        year  = today.year + ((today.month - i - 1) // 12)
        count = documents.filter(
            upload_date__year=year,
            upload_date__month=month,
        ).count()
        monthly_data.append({
            'label': ['Jan','Fév','Mar','Avr','Mai','Jun','Jul','Aoû','Sep','Oct','Nov','Déc'][month - 1],
            'count': count,
        })
    max_monthly = max((m['count'] for m in monthly_data), default=1) or 1

    # ── Conformité par norme ─────────────────────────────────
    norme_stats = defaultdict(lambda: {'total': 0, 'conformes': 0})
    for doc in documents.filter(status='termine'):
        n = doc.norme or ('Carbone' if doc.type_analyse == 'carbone' else '—')
        norme_stats[n]['total'] += 1
        if doc.is_conform is True:
            norme_stats[n]['conformes'] += 1
    norme_conformite = []
    for norme_name, vals in norme_stats.items():
        pct = round(vals['conformes'] / vals['total'] * 100) if vals['total'] else 0
        norme_conformite.append({'norme': norme_name, 'total': vals['total'], 'conformes': vals['conformes'], 'pct': pct})
    norme_conformite.sort(key=lambda x: -x['total'])

    # ── Terminés ce mois ─────────────────────────────────────
    termine_ce_mois = documents.filter(
        status='termine',
        upload_date__year=today.year,
        upload_date__month=today.month,
    ).count()

    context = {
        'documents':   documents,
        'energie_docs': energie_docs,
        'carbone_docs': carbone_docs,

        'total_projects':  total_projects,
        'compliance_rate': compliance_rate,
        'pending_count':   pending_count,
        'old_pending':     old_pending,
        'recent_devis':    recent_devis,
        'devis_en_attente': devis_en_attente,
        # Détail par statut
        'docs_energie_recu':     energie_docs.filter(status='recu'),
        'docs_energie_en_cours': energie_docs.filter(status='en_cours'),
        'docs_energie_termine':  energie_docs.filter(status='termine'),
        'docs_carbone_recu':     carbone_docs.filter(status='recu'),
        'docs_carbone_en_cours': carbone_docs.filter(status='en_cours'),
        'docs_carbone_termine':  carbone_docs.filter(status='termine'),

        # Compteurs
        'count_energie': energie_docs.count(),
        'count_carbone': carbone_docs.count(),

        # Nouveaux : stats enrichies
        'monthly_data':      monthly_data,
        'max_monthly':       max_monthly,
        'norme_conformite':  norme_conformite,
        'termine_ce_mois':   termine_ce_mois,
    }
    return render(request, 'main/home.html', context)


@login_required(login_url='/login/')
def results(request):
    documents = Document.objects.filter(is_active=True)
    context = {
        'documents': documents,
        're2020_requirements': fetch_re2020_requirements(),
        'rt2012_requirements': fetch_rt2012_requirements(),
    }
    return render(request, 'main/results.html', context)


@login_required(login_url='/login/')
def history(request):
    qs = Document.objects.filter(is_active=True).order_by('-upload_date')

    # ── Filtres serveur-side (couvrent toutes les pages) ──
    q = request.GET.get('q', '').strip()
    filtre_status = request.GET.get('status', '')
    filtre_type   = request.GET.get('type', '')

    if q:
        from django.db.models import Q
        qs = qs.filter(
            Q(name__icontains=q) |
            Q(client_name__icontains=q) |
            Q(client_email__icontains=q)
        )
    if filtre_status:
        qs = qs.filter(status=filtre_status)
    if filtre_type:
        qs = qs.filter(type_analyse=filtre_type)

    paginator = Paginator(qs, 20)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    return render(request, 'main/history.html', {
        'documents':      page_obj,
        'paginator':      paginator,
        'page_obj':       page_obj,
        'q':              q,
        'filtre_status':  filtre_status,
        'filtre_type':    filtre_type,
    })


@login_required(login_url='/login/')
def export_csv_history(request):
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="dossiers_conformexpert.csv"'
    response.write('\ufeff')
    writer = csv.writer(response, delimiter=';')
    writer.writerow(['Référence', 'Nom du dossier', 'Client', 'Email client', 'Type', 'Norme', 'Statut', 'Conformité', 'Date de dépôt'])
    for doc in Document.objects.all().order_by('-upload_date'):
        conform = '—'
        if doc.is_conform is True:
            conform = 'Conforme'
        elif doc.is_conform is False:
            conform = 'Non conforme'
        writer.writerow([
            f'DOC-{doc.id:04d}',
            doc.name,
            doc.client_name or '—',
            doc.client_email or '—',
            'Bilan carbone' if doc.type_analyse == 'carbone' else 'Validation thermique',
            doc.norme or '—',
            {'recu': 'Reçu', 'en_cours': 'En cours', 'termine': 'Terminé'}.get(doc.status, doc.status),
            conform,
            doc.upload_date.strftime('%d/%m/%Y') if doc.upload_date else '—',
        ])
    return response


@login_required(login_url='/login/')
def ia_rapport_status(request, doc_id):
    doc = get_object_or_404(Document, id=doc_id)
    return JsonResponse({
        'has_rapport': bool(doc.rapport_ia_json),
        'status': doc.status,
    })


@login_required(login_url='/login/')
def settings_view(request):
    from main.templatetags.conformity_tags import get_seuils, NORME_FIELDS, NORMES_PAR_PAYS

    PAYS_LABELS = {
        'FR': '🇫🇷 France',
        'BE': '🇧🇪 Belgique',
        'CH': '🇨🇭 Suisse',
        'CA': '🇨🇦 Canada',
        'LU': '🇱🇺 Luxembourg',
    }
    NORME_LABELS = {
        'RT2012': 'RT 2012', 'RE2020': 'RE 2020', 'PEB': 'PEB',
        'MINERGIE': 'Minergie', 'SIA380': 'SIA 380',
        'CNEB2015': 'CNEB 2015', 'CNEB2020': 'CNEB 2020', 'LENOZ': 'LENOZ',
    }

    seuils_par_pays = []
    for pays_code, normes in NORMES_PAR_PAYS.items():
        normes_data = []
        for norme_code in normes:
            seuils = get_seuils('maison', 'H2', pays_code, norme_code)
            fields = NORME_FIELDS.get(norme_code, [])
            fields_seuils = []
            for field, label, unit in fields:
                val = seuils.get(field, '—')
                if isinstance(val, float) and val == int(val):
                    val = int(val)
                fields_seuils.append((label, val, unit))
            normes_data.append((norme_code, NORME_LABELS.get(norme_code, norme_code), fields_seuils))
        seuils_par_pays.append((pays_code, PAYS_LABELS.get(pays_code, pays_code), normes_data))

    carbone_seuils = [
        ("Ic énergie RE2020 — seuil max (maison individuelle)",  4,    "kgCO2eq/m².an"),
        ("Ic énergie RE2020 — seuil max (logement collectif)",   6.5,  "kgCO2eq/m².an"),
        ("Ic construction RE2020 — seuil max (mi-vie 2025)",   640,    "kgCO2eq/m²"),
        ("Ic construction RE2020 — seuil max (mi-vie 2028)",   590,    "kgCO2eq/m²"),
        ("Réduction conso. décret tertiaire 2030",              40,    "%"),
        ("Réduction conso. décret tertiaire 2040",              50,    "%"),
        ("Réduction conso. décret tertiaire 2050",              60,    "%"),
        ("Émissions GES DPE — seuil passoire (classe F)",      70,    "kgCO2eq/m².an"),
        ("Émissions GES DPE — seuil passoire (classe G)",     100,    "kgCO2eq/m².an"),
    ]

    return render(request, 'main/settings.html', {
        'seuils_par_pays': seuils_par_pays,
        'carbone_seuils': carbone_seuils,
    })


def update_re2020(request):
    if request.method == 'POST':
        messages.success(request, 'Paramètres RE2020 mis à jour.')
    else:
        messages.error(request, 'Méthode invalide.')
    return redirect('settings')


# ──────────────────────────────────────────────────────────────
# ÉDITION DOCUMENT
# ──────────────────────────────────────────────────────────────

# Constante partagée : champs éditables par norme
ALL_NORME_FIELDS = {
    'RT2012': [
        ('rt2012_bbio',         'Bbio',       ''),
        ('rt2012_cep',          'Cep',        'kWh ep/m².an'),
        ('rt2012_tic',          'Tic',        '°C'),
        ('rt2012_airtightness', 'Étanchéité', 'm³/h.m²'),
        ('rt2012_enr',          'ENR',        ''),
    ],
    'RE2020': [
        ('re2020_energy_efficiency', 'Cep,nr',          'kWh/m².an'),
        ('re2020_carbon_emissions',  'Ic énergie CO₂', 'kgCO2eq/m².an'),
        ('re2020_thermal_comfort',   'DH (confort été)', 'DH'),
    ],
    'PEB': [
        ('peb_espec',     'Espec',      'kWh/m².an'),
        ('peb_ew',        'Ew',         ''),
        ('peb_u_mur',     'U mur',      'W/m².K'),
        ('peb_u_toit',    'U toit',     'W/m².K'),
        ('peb_u_plancher','U plancher', 'W/m².K'),
    ],
    'MINERGIE': [
        ('minergie_qh',   'Qh',   'kWh/m².an'),
        ('minergie_qtot', 'Qtot', 'kWh/m².an'),
        ('minergie_n50',  'n50',  'h⁻¹'),
    ],
    'SIA380': [
        ('sia380_qh', 'Qh', 'kWh/m².an'),
    ],
    'CNEB2015': [
        ('cneb_ei',          'Intensité énergétique', 'kWh/m².an'),
        ('cneb_u_mur',       'U mur',                 'W/m².K'),
        ('cneb_u_toit',      'U toit',                'W/m².K'),
        ('cneb_u_fenetre',   'U fenêtre',             'W/m².K'),
        ('cneb_infiltration','Infiltration',           'L/s.m²'),
    ],
    'CNEB2020': [
        ('cneb_ei',          'Intensité énergétique', 'kWh/m².an'),
        ('cneb_u_mur',       'U mur',                 'W/m².K'),
        ('cneb_u_toit',      'U toit',                'W/m².K'),
        ('cneb_u_fenetre',   'U fenêtre',             'W/m².K'),
        ('cneb_infiltration','Infiltration',           'L/s.m²'),
    ],
    'LENOZ': [
        ('lenoz_ep',    'Énergie primaire', 'kWh/m².an'),
        ('lenoz_ew',    'Ew',               ''),
        ('lenoz_u_mur', 'U mur',            'W/m².K'),
        ('lenoz_u_toit','U toit',           'W/m².K'),
    ],
}

STATUS_CHOICES = [
    ('recu',     'Dossier reçu'),
    ('en_cours', 'Analyse en cours'),
    ('termine',  'Analyse terminée'),
]


@login_required(login_url='/login/')
def edit_document(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    norme_fields = ALL_NORME_FIELDS.get(document.norme, [])

    if request.method == 'POST':

        # ── Mise à jour du statut ──────────────────────────────
        new_status = request.POST.get('status')
        old_status = document.status
        if new_status in dict(STATUS_CHOICES):
            document.status = new_status

        # ── Mise à jour de la norme ────────────────────────────
        new_norme = request.POST.get('norme', document.norme)
        if new_norme in ALL_NORME_FIELDS:
            document.norme = new_norme

        # ── Champs thermiques de toutes les normes ─────────────
        for fields in ALL_NORME_FIELDS.values():
            for field, _, _ in fields:
                val = request.POST.get(field, '').strip()
                if val:
                    try:
                        setattr(document, field, float(val.replace(',', '.')))
                    except ValueError:
                        pass

        # ── Infos client ───────────────────────────────────────
        document.client_name  = request.POST.get('client_name', '').strip()
        document.client_email = request.POST.get('client_email', '').strip()
        document.admin_notes  = request.POST.get('admin_notes', '').strip()
        document.save()

        # ── Emails selon changement de statut ──────────────────
        if old_status != new_status:
            if new_status == 'recu':
                try:
                    devis = document.devis.filter(statut='en_attente').first()
                except Exception:
                    devis = None
                send_mail_validation_devis(document, devis)
                messages.info(
                    request,
                    f'Email de validation + devis envoyé à {document.client_email}.'
                    if document.client_email else "Pas d'email client renseigné."
                )
            elif new_status == 'en_cours':
                send_mail_analyse_commence(document)
                messages.info(
                    request,
                    f'Email "analyse commencée" envoyé à {document.client_email}.'
                    if document.client_email else "Pas d'email client renseigné."
                )
            elif new_status == 'termine':
                send_mail_analyse_terminee(document)
                messages.info(
                    request,
                    f'Email "rapport disponible" envoyé à {document.client_email}.'
                    if document.client_email else "Pas d'email client renseigné."
                )

        messages.success(request, f'Dossier « {document.name} » mis à jour.')
        return redirect('edit_document', doc_id=doc_id)

    # ── Pré-parser le rapport IA pour rendu serveur-side ──
    rapport_ia = None
    rapport_ia_verdict = None
    rapport_ia_score   = None
    rapport_ia_resume  = None
    rapport_ia_fiabilite = None
    if document.rapport_ia_json:
        try:
            rapport_ia = json.loads(document.rapport_ia_json)
            rapport_ia_verdict   = rapport_ia.get('verdict') or rapport_ia.get('verdict_technique') or rapport_ia.get('verdict_energie')
            rapport_ia_score     = rapport_ia.get('score_global')
            rapport_ia_resume    = rapport_ia.get('resume_executif') or ''
            rapport_ia_fiabilite = rapport_ia.get('fiabilite_rapport')
        except Exception:
            rapport_ia = None

    return render(request, 'main/edit_document.html', {
        'document':        document,
        'status_choices':  STATUS_CHOICES,
        'rt2012_fields':   ALL_NORME_FIELDS.get('RT2012', []),
        're2020_fields':   ALL_NORME_FIELDS.get('RE2020', []),
        'norme_fields':    norme_fields,
        'all_norme_fields': ALL_NORME_FIELDS,
        'norme_choices':   Document.NORME_CHOICES,
        'rapport_ia':          rapport_ia,
        'rapport_ia_verdict':  rapport_ia_verdict,
        'rapport_ia_score':    rapport_ia_score,
        'rapport_ia_resume':   rapport_ia_resume,
        'rapport_ia_fiabilite':rapport_ia_fiabilite,
        'email_steps': [
            ('1', '#60a5fa', 'rgba(59,130,246,.12)',  'Confirmation réception',       'Confirmer la réception du dossier',      'reception'),
            ('2', '#c8a84b', 'rgba(200,168,75,.12)',  'Envoi du devis',               "Devis avec bouton d'acceptation",        'devis'),
            ('3', '#2dd4bf', 'rgba(20,184,166,.12)',  'Début analyse + lien suivi',   "Notifier le démarrage de l'analyse",     'analyse_commence'),
            ('4', '#27c93f', 'rgba(39,201,63,.12)',   'Rapport final disponible',     'Rapport téléchargeable sur le lien suivi', 'analyse_terminee'),
        ],
    })


@login_required(login_url='/login/')
def send_email_manual(request, doc_id, email_type):
    if request.method != 'POST':
        return redirect('edit_document', doc_id=doc_id)

    document = get_object_or_404(Document, id=doc_id)
    if not document.client_email:
        messages.error(request, 'Aucun email client renseigné.')
        return redirect('edit_document', doc_id=doc_id)

    if email_type == 'reception':
        send_mail_reception(document)
        messages.success(request, f'Email de réception envoyé à {document.client_email}.')

    elif email_type == 'devis':
        if request.POST.get('create_devis'):
            montant = request.POST.get('montant', '').strip()
            devis = Devis(
                client_nom=document.client_name or document.client_email,
                client_email=document.client_email,
                projet_nom=request.POST.get('projet_nom', document.name),
                norme=request.POST.get('norme', 'RE2020'),
                montant=float(montant) if montant else None,
                notes=request.POST.get('notes', ''),
                statut='en_attente',
                document=document,
            )
            devis.save()
        else:
            devis = document.devis.filter(statut='en_attente').first()
        send_mail_validation_devis(document, devis)
        messages.success(request, f'Email devis envoyé à {document.client_email}.')

    elif email_type == 'analyse_commence':
        send_mail_analyse_commence(document)
        messages.success(request, f'Email "analyse démarrée" envoyé à {document.client_email}.')

    elif email_type == 'analyse_terminee':
        send_mail_analyse_terminee(document)
        messages.success(request, f'Email "rapport disponible" envoyé à {document.client_email}.')

    return redirect('edit_document', doc_id=doc_id)


@login_required(login_url='/login/')
def upload_rapport_pdf(request, doc_id):
    if request.method == 'POST' and request.FILES.get('rapport_pdf'):
        document = get_object_or_404(Document, id=doc_id)
        document.rapport_pdf = request.FILES['rapport_pdf']
        document.save()
        messages.success(request, 'Rapport PDF uploadé avec succès.')
    return redirect('edit_document', doc_id=doc_id)


@login_required(login_url='/login/')
def delete_document(request, doc_id):
    if request.method == 'POST':
        document = get_object_or_404(Document, id=doc_id)
        document.delete()
        messages.success(request, 'Dossier supprimé.')
    return redirect('history')


# ──────────────────────────────────────────────────────────────
# VÉRIFICATION DES SEUILS VIA IA
# ──────────────────────────────────────────────────────────────

@csrf_exempt
def verifier_seuils(request):
    import os
    import urllib.request
    import urllib.error

    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Session expirée — veuillez vous reconnecter'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode invalide'}, status=405)

    norme = request.POST.get('norme', 'RE2020')
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

    if not ANTHROPIC_API_KEY:
        return JsonResponse(
            {'error': 'Clé API Anthropic manquante — ajoutez ANTHROPIC_API_KEY dans vos variables Railway'},
            status=500,
        )

    NORME_CONTEXTE = {
        'RT2012':   'la réglementation thermique RT 2012 française (arrêté du 26 octobre 2010)',
        'RE2020':   'la réglementation environnementale RE 2020 française (décret du 29 juillet 2021)',
        'PEB':      'la performance énergétique des bâtiments PEB en Belgique',
        'MINERGIE': 'le label Minergie en Suisse',
        'SIA380':   'la norme SIA 380/1 suisse',
        'CNEB2015': "le Code National de l'Énergie pour les Bâtiments CNEB 2015 au Canada",
        'CNEB2020': "le Code National de l'Énergie pour les Bâtiments CNEB 2020 au Canada",
        'LENOZ':    'le label LENOZ au Luxembourg',
    }
    VALEURS_ACTUELLES = {
        'RT2012':   {'Bbio max (maison)': 60, 'Bbio max (collectif)': 80, 'Cep max': 50,
                     'Tic max H2': 27, 'Étanchéité maison': 0.6, 'ENR min': 1.0},
        'RE2020':   {'Cep,nr max (maison)': 100, 'DH max H2': 1250,
                     'Ic énergie max': 160, 'Ic construction max': 640},
        'PEB':      {'Espec max': 100, 'U mur max': 0.24, 'U toit max': 0.20, 'U plancher max': 0.30},
        'MINERGIE': {'Qh max (maison)': 60, 'Qtot max': 38, 'n50 max': 0.6},
        'SIA380':   {'Qh max (maison)': 90},
        'CNEB2015': {'EI max (maison)': 170, 'U mur max': 0.24, 'U toit max': 0.18,
                     'U fenêtre max': 1.8, 'Infiltration max': 0.30},
        'CNEB2020': {'EI max (maison)': 150, 'U mur max': 0.21, 'U toit max': 0.16,
                     'U fenêtre max': 1.6, 'Infiltration max': 0.25},
        'LENOZ':    {'EP max (maison)': 90, 'Ew max': 100, 'U mur max': 0.22, 'U toit max': 0.17},
    }

    valeurs  = VALEURS_ACTUELLES.get(norme, {})
    contexte = NORME_CONTEXTE.get(norme, norme)

    prompt = f"""Tu es un expert en réglementation thermique et énergétique des bâtiments.

Vérifie si les seuils suivants pour {contexte} sont toujours officiellement valides en {__import__('datetime').date.today().year}.

Valeurs actuellement dans notre système :
{json.dumps(valeurs, ensure_ascii=False, indent=2)}

Réponds UNIQUEMENT en JSON valide avec cette structure exacte, sans markdown ni explication :
{{
  "a_jour": true,
  "date_derniere_mise_a_jour": "mois et année",
  "modifications": [
    {{
      "critere": "nom du critère",
      "valeur_actuelle": 100,
      "valeur_officielle": 90,
      "commentaire": "explication courte"
    }}
  ],
  "resume": "phrase courte résumant le statut",
  "source": "texte officiel de référence"
}}

Si tout est à jour, "modifications" sera [] et "a_jour" sera true."""

    try:
        payload = json.dumps({
            "model": "claude-sonnet-4-5",
            "max_tokens": 1500,
            "messages": [{"role": "user", "content": prompt}],
        }).encode('utf-8')

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            raw = result['content'][0]['text'].strip().replace('```json', '').replace('```', '').strip()
            return JsonResponse({'success': True, 'norme': norme, 'resultat': json.loads(raw)})

    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8')
        print(f"ANTHROPIC API ERROR {e.code}: {body}")
        return JsonResponse({'error': f'API {e.code}: {body}'}, status=500)
    except Exception as e:
        print(f"VERIFIER_SEUILS ERROR: {e}")
        return JsonResponse({'error': str(e)}, status=500)


# ──────────────────────────────────────────────────────────────
# RAPPORT PDF (ReportLab)
# ──────────────────────────────────────────────────────────────

def download_rapport_word(request, doc_id):
    from docx import Document as DocxDocument
    from io import BytesIO

    document = get_object_or_404(Document, id=doc_id)
    doc = DocxDocument()
    doc.add_heading(f'Rapport ConformExpert — {document.name}', 0)
    doc.add_paragraph(f'Référence : DOC-{document.id:04d}')
    doc.add_paragraph(f'Client : {document.client_name or "—"}')
    doc.add_paragraph(f'Date : {document.upload_date.strftime("%d/%m/%Y")}')

    doc.add_heading('RT2012', level=1)
    for label, val in [
        ('Bbio', document.rt2012_bbio), ('Cep', document.rt2012_cep),
        ('Tic', document.rt2012_tic), ('Étanchéité', document.rt2012_airtightness),
        ('ENR', document.rt2012_enr),
    ]:
        doc.add_paragraph(f'{label} : {val if val is not None else "—"}')

    doc.add_heading('RE2020', level=1)
    for label, val in [
        ('Cep,nr', document.re2020_energy_efficiency),
        ('Ic énergie', document.re2020_carbon_emissions),
        ('DH', document.re2020_thermal_comfort),
    ]:
        doc.add_paragraph(f'{label} : {val if val is not None else "—"}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = document.name.replace(' ', '_')
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="rapport_{safe_name}.docx"'
    return response


def download_report(request, document_id):
    """Génère et renvoie le rapport PDF ReportLab complet."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from main.templatetags.conformity_tags import get_seuils, CRITERIA_GREATER_EQUAL

    document = get_object_or_404(Document, id=document_id)

    PAGE_W, PAGE_H = A4
    ML = MR = MT = 2 * cm
    MB = 2.5 * cm
    W  = PAGE_W - ML - MR

    # ── Palette de couleurs ──
    NAVY  = colors.HexColor('#0C1929')
    GOLD  = colors.HexColor('#C8A84B')
    GOLD_L = colors.HexColor('#F5EDD0')
    GREEN  = colors.HexColor('#1A9E2E')
    GREEN_L = colors.HexColor('#E8F8EE')
    RED   = colors.HexColor('#C62828')
    RED_L = colors.HexColor('#FEF0F0')
    LGRAY = colors.HexColor('#F8F8FC')
    MGRAY = colors.HexColor('#E0E0E8')
    WHITE = colors.white
    MUTED = colors.HexColor('#888899')
    TEXT  = colors.HexColor('#1A1A2E')

    def st(name, **kw):
        d = dict(fontName='Helvetica', fontSize=9, textColor=TEXT, leading=14, spaceAfter=0)
        d.update(kw)
        return ParagraphStyle(name, **d)

    seuils    = get_seuils(document.building_type, document.climate_zone, document.pays, document.norme)
    is_conform = document.is_conform
    norme     = document.norme
    pays_map  = {'FR': 'France', 'BE': 'Belgique', 'CH': 'Suisse', 'CA': 'Canada', 'LU': 'Luxembourg'}
    pays_label = pays_map.get(document.pays, document.pays)
    today_str  = date.today().strftime("%d/%m/%Y")

    if is_conform is True:
        verdict_txt, verdict_col, verdict_bg = "Conforme", GREEN, GREEN_L
    elif is_conform is False:
        verdict_txt, verdict_col, verdict_bg = "Non Conforme", RED, RED_L
    else:
        verdict_txt, verdict_col, verdict_bg = "En cours d'analyse", MUTED, LGRAY

    # ── Helpers ──
    def section_title(num, title):
        label = f"{num}.  {title.upper()}"
        return [
            HRFlowable(width=W, thickness=1.5, color=GOLD, spaceAfter=5, spaceBefore=10),
            Paragraph(label, st('sh', fontName='Helvetica-Bold', fontSize=8,
                                textColor=GOLD, characterSpacing=0.8, spaceAfter=8)),
        ]

    def info_table(rows):
        data = [
            [Paragraph(k, st('ik', fontSize=8, textColor=MUTED)),
             Paragraph(v, st('iv', fontName='Helvetica-Bold', fontSize=9))]
            for k, v in rows
        ]
        t = Table(data, colWidths=[4.5 * cm, W - 4.5 * cm])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (0, -1), LGRAY),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING',   (0, 0), (-1, -1), 8),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 8),
            ('LINEBELOW',     (0, 0), (-1, -2), 0.5, MGRAY),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return t

    def criteria_section(title, rows_data):
        rows = [r for r in rows_data if r is not None]
        if not rows:
            return []
        header = [
            Paragraph("Critère",  st('th',  fontName='Helvetica-Bold', fontSize=8, textColor=WHITE)),
            Paragraph("Valeur",   st('th2', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE, alignment=TA_CENTER)),
            Paragraph("Seuil",    st('th3', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE, alignment=TA_CENTER)),
            Paragraph("Unité",    st('th4', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE, alignment=TA_CENTER)),
            Paragraph("Résultat", st('th5', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE, alignment=TA_CENTER)),
        ]
        data  = [header] + rows
        col_w = [7 * cm, 2.2 * cm, 2.2 * cm, 2.3 * cm, 3.3 * cm]
        t = Table(data, colWidths=col_w, repeatRows=1)
        style = [
            ('BACKGROUND',    (0, 0), (-1, 0), NAVY),
            ('TOPPADDING',    (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING',   (0, 0), (-1, -1), 6),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
            ('LINEBELOW',     (0, 1), (-1, -1), 0.5, MGRAY),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ]
        for i in range(2, len(data), 2):
            style.append(('BACKGROUND', (0, i), (-1, i), LGRAY))
        t.setStyle(TableStyle(style))
        return [t, Spacer(1, 0.4 * cm)]

    def criteria_row(label, value, key, unit=""):
        if value is None:
            return None
        limit = seuils.get(key, "—")
        sign  = ">=" if key in CRITERIA_GREATER_EQUAL else "<="
        if isinstance(limit, (int, float)):
            ok = value >= limit if key in CRITERIA_GREATER_EQUAL else value <= limit
        else:
            ok = False
        res_style = st('res', fontName='Helvetica-Bold', fontSize=8,
                       textColor=GREEN if ok else RED, alignment=TA_CENTER)
        return [
            Paragraph(label, st('cl', fontSize=8.5)),
            Paragraph(f"<b>{value}</b>", st('cv', fontName='Helvetica-Bold', fontSize=9, alignment=TA_CENTER)),
            Paragraph(f"{sign} {limit}", st('cs', fontSize=8, textColor=MUTED, alignment=TA_CENTER)),
            Paragraph(unit, st('cu', fontSize=8, textColor=MUTED, alignment=TA_CENTER)),
            Paragraph("Conforme" if ok else "Non conforme", res_style),
        ]

    def reco_block(prefix, title, text, bg, left_color):
        full_title = f"<b>{prefix}  {title}</b>"
        data = [
            [Paragraph(full_title, st('rbt', fontName='Helvetica-Bold', fontSize=9, textColor=TEXT, spaceAfter=3))],
            [Paragraph(text, st('rbd', fontSize=8.5, textColor=colors.HexColor('#444455'), leading=13))],
        ]
        t = Table(data, colWidths=[W])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), bg),
            ('LINEBEFORE',    (0, 0), (0, -1), 3, left_color),
            ('TOPPADDING',    (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING',   (0, 0), (-1, -1), 12),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 12),
        ]))
        return [t, Spacer(1, 0.25 * cm)]

    # ── Page de couverture (canvas) ──
    def draw_cover(c, doc_obj):
        c.saveState()
        c.setFillColor(NAVY)
        c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)

        c.setFillColor(colors.HexColor('#C8A84B11'))
        c.circle(PAGE_W, PAGE_H, 180, fill=1, stroke=0)
        c.setFillColor(colors.HexColor('#C8A84B0A'))
        c.circle(0, 0, 130, fill=1, stroke=0)

        c.setFont('Helvetica-Bold', 22)
        c.setFillColor(WHITE)
        c.drawString(ML, PAGE_H - 3.5 * cm, "Conform")
        c.setFillColor(GOLD)
        lw = c.stringWidth("Conform", 'Helvetica-Bold', 22)
        c.drawString(ML + lw, PAGE_H - 3.5 * cm, "Expert")

        c.setFont('Helvetica-Bold', 7.5)
        c.setFillColor(GOLD)
        c.drawString(ML, PAGE_H - 5 * cm, "RAPPORT D'ANALYSE DE CONFORMITÉ THERMIQUE")
        c.setStrokeColor(GOLD)
        c.setLineWidth(1)
        c.line(ML, PAGE_H - 5.3 * cm, ML + 6 * cm, PAGE_H - 5.3 * cm)

        title = document.name
        c.setFont('Helvetica-Bold', 24)
        c.setFillColor(WHITE)
        max_w = PAGE_W - ML - MR - 1 * cm
        while c.stringWidth(title, 'Helvetica-Bold', 24) > max_w and len(title) > 10:
            title = title[:-1]
        if title != document.name:
            title = title[:-3] + "..."
        c.drawString(ML, PAGE_H - 7 * cm, title)

        c.setFont('Helvetica', 11)
        c.setFillColor(colors.HexColor('#AAAACC'))
        subtitle = f"{document.get_building_type_display()}  ·  {norme}  ·  {pays_label}"
        c.drawString(ML, PAGE_H - 8.2 * cm, subtitle)

        # Verdict
        v_y, v_x, v_w, v_h = PAGE_H - 10.5 * cm, ML, 9 * cm, 1.5 * cm
        c.setFillColor(verdict_bg)
        c.roundRect(v_x, v_y, v_w, v_h, 20, fill=1, stroke=0)
        c.setStrokeColor(verdict_col)
        c.setLineWidth(1.5)
        c.roundRect(v_x, v_y, v_w, v_h, 20, fill=0, stroke=1)
        c.setFont('Helvetica-Bold', 12)
        c.setFillColor(verdict_col)
        tw = c.stringWidth(verdict_txt, 'Helvetica-Bold', 12)
        c.drawString(v_x + (v_w - tw) / 2, v_y + 0.45 * cm, verdict_txt)

        c.setStrokeColor(colors.HexColor('#C8A84B44'))
        c.setLineWidth(0.5)
        c.line(ML, PAGE_H - 12.5 * cm, PAGE_W - MR, PAGE_H - 12.5 * cm)

        meta = [
            ("REFERENCE",        f"DOC-{document.id:04d}"),
            ("DATE DU RAPPORT",  today_str),
            ("CLIENT",           document.client_name or "—"),
            ("NORME",            norme),
            ("TYPE DE BÂTIMENT", document.get_building_type_display()),
            ("PAYS",             pays_label),
        ]
        cols   = 3
        col_w2 = (PAGE_W - ML - MR) / cols
        for i, (lbl, val) in enumerate(meta):
            col = i % cols
            row = i // cols
            x = ML + col * col_w2
            y = PAGE_H - 13.5 * cm - row * 1.8 * cm
            c.setFont('Helvetica', 7)
            c.setFillColor(MUTED)
            c.drawString(x, y, lbl)
            c.setFont('Helvetica-Bold', 10)
            c.setFillColor(WHITE)
            c.drawString(x, y - 0.5 * cm, val)

        c.setFillColor(colors.HexColor('#C8A84B33'))
        c.setStrokeColor(colors.HexColor('#00000000'))
        c.rect(0, 0, PAGE_W, 1.8 * cm, fill=1, stroke=0)
        c.setFont('Helvetica', 8)
        c.setFillColor(colors.HexColor('#666677'))
        c.drawString(ML, 0.65 * cm, "ConformExpert  ·  Analyse documentaire indépendante")
        c.setFillColor(GOLD)
        txt_r = "Confidentiel  ·  Usage interne"
        tw2 = c.stringWidth(txt_r, 'Helvetica', 8)
        c.drawString(PAGE_W - MR - tw2, 0.65 * cm, txt_r)
        c.restoreState()

    def draw_page(c, doc_obj):
        """Footer sur les pages intérieures."""
        c.saveState()
        c.setStrokeColor(MGRAY)
        c.setLineWidth(0.5)
        c.line(ML, MB - 0.5 * cm, PAGE_W - MR, MB - 0.5 * cm)
        c.setFont('Helvetica', 7.5)
        c.setFillColor(MUTED)
        c.drawString(ML, MB - 1 * cm, f"ConformExpert  ·  Analyse indépendante {norme}  ·  {pays_label}")
        page_num = str(doc_obj.page)
        tw = c.stringWidth(f"Page {page_num}", 'Helvetica', 7.5)
        c.drawString(PAGE_W - MR - tw, MB - 1 * cm, f"Page {page_num}")
        c.restoreState()

    # ── Story ──
    story = []
    story.append(Spacer(1, PAGE_H - MT - MB))
    story.append(PageBreak())

    # Sommaire
    story += section_title("", "Sommaire")
    story.append(Paragraph(
        f"Rapport d'analyse — {document.name}",
        st('ts', fontSize=8.5, textColor=MUTED, spaceAfter=12),
    ))
    toc_items = [
        ("1.  Résumé exécutif & verdict global",             "3"),
        ("2.  Informations du dossier",                      "3"),
        (f"3.  Analyse {norme} — Critères de conformité",   "4"),
        ("4.  Recommandations & points d'attention",         "5"),
    ]
    if document.admin_notes:
        toc_items.append(("5.  Notes & observations de l'expert", "5"))
    toc_items.append(("6.  Mentions légales & disclaimer", "6"))

    for label, pg in toc_items:
        row = Table([[
            Paragraph(label, st('tl', fontSize=10)),
            Paragraph(pg, st('tp', fontName='Helvetica-Bold', fontSize=9, textColor=GOLD, alignment=TA_RIGHT)),
        ]], colWidths=[W - 1.5 * cm, 1.5 * cm])
        row.setStyle(TableStyle([
            ('TOPPADDING',    (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
            ('LINEBELOW',     (0, 0), (-1, -1), 0.5, MGRAY),
        ]))
        story.append(row)

    story.append(PageBreak())

    # Résumé + verdict
    story += section_title("1", "Résumé exécutif & verdict global")
    vb_data = [[
        Paragraph(f"VERDICT  —  {norme}",
                  st('vbl', fontName='Helvetica-Bold', fontSize=8, textColor=GOLD,
                     characterSpacing=0.8, spaceAfter=4)),
        Paragraph(verdict_txt,
                  st('vbv', fontName='Helvetica-Bold', fontSize=16,
                     textColor=verdict_col, alignment=TA_RIGHT)),
    ]]
    vb = Table(vb_data, colWidths=[W * 0.55, W * 0.45])
    vb.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), NAVY),
        ('TOPPADDING',    (0, 0), (-1, -1), 14),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 14),
        ('LEFTPADDING',   (0, 0), (-1, -1), 16),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 16),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(vb)
    story.append(Spacer(1, 0.6 * cm))

    # Infos dossier
    story += section_title("2", "Informations du dossier")
    story.append(info_table([
        ("Référence",        f"DOC-{document.id:04d}"),
        ("Norme analysée",   norme),
        ("Pays",             pays_label),
        ("Type de bâtiment", document.get_building_type_display()),
        ("Zone climatique",  f"Zone {document.climate_zone}" if document.climate_zone else "—"),
        ("Date de dépôt",    document.upload_date.strftime("%d/%m/%Y")),
        ("Date du rapport",  today_str),
        ("Client",           document.client_name or "—"),
        ("Email client",     document.client_email or "—"),
    ]))
    story.append(PageBreak())

    # Critères
    story += section_title("3", f"{norme} — Critères de conformité")
    if norme == 'RT2012':
        story += criteria_section("RT2012", [
            criteria_row("Bbio — Besoins bioclimatiques",        document.rt2012_bbio,         "rt2012_bbio"),
            criteria_row("Cep — Consommation énergie primaire",  document.rt2012_cep,           "rt2012_cep",          "kWh ep/m².an"),
            criteria_row("Tic — Température intérieure conv.",   document.rt2012_tic,           "rt2012_tic",          "°C"),
            criteria_row("Étanchéité à l'air",                   document.rt2012_airtightness,  "rt2012_airtightness", "m³/h.m²"),
            criteria_row("ENR — Énergies renouvelables",         document.rt2012_enr,           "rt2012_enr"),
        ])
    elif norme == 'RE2020':
        story += criteria_section("RE2020", [
            criteria_row("Cep,nr — Énergie non renouvelable",       document.re2020_energy_efficiency, "re2020_energy_efficiency", "kWh/m².an"),
            criteria_row("Ic énergie — Émissions CO₂ exploitation", document.re2020_carbon_emissions,  "re2020_carbon_emissions",  "kgCO₂/m².an"),
            criteria_row("DH — Degrés-heures (confort été)",        document.re2020_thermal_comfort,   "re2020_thermal_comfort",   "DH"),
        ])
    elif norme == 'PEB':
        story += criteria_section("PEB", [
            criteria_row("Espec — Énergie spécifique",            document.peb_espec,      "peb_espec",      "kWh/m².an"),
            criteria_row("Ew — Indicateur global de performance", document.peb_ew,         "peb_ew"),
            criteria_row("U mur",                                 document.peb_u_mur,      "peb_u_mur",      "W/m².K"),
            criteria_row("U toit",                                document.peb_u_toit,     "peb_u_toit",     "W/m².K"),
            criteria_row("U plancher",                            document.peb_u_plancher, "peb_u_plancher", "W/m².K"),
        ])
    elif norme == 'MINERGIE':
        story += criteria_section("MINERGIE", [
            criteria_row("Qh — Chaleur de chauffage annuelle",  document.minergie_qh,   "minergie_qh",   "kWh/m².an"),
            criteria_row("Qtot — Énergie totale pondérée",      document.minergie_qtot, "minergie_qtot", "kWh/m².an"),
            criteria_row("n50 — Taux de renouvellement d'air",  document.minergie_n50,  "minergie_n50",  "h⁻¹"),
        ])
    elif norme == 'SIA380':
        story += criteria_section("SIA380", [
            criteria_row("Qh — Chaleur de chauffage (SIA 380/1)", document.sia380_qh, "sia380_qh", "kWh/m².an"),
        ])
    elif norme in ('CNEB2015', 'CNEB2020'):
        story += criteria_section(norme, [
            criteria_row("Intensité énergétique",               document.cneb_ei,           "cneb_ei",           "kWh/m².an"),
            criteria_row("U mur — Valeur thermique enveloppe",  document.cneb_u_mur,        "cneb_u_mur",        "W/m².K"),
            criteria_row("U toit — Valeur thermique toiture",   document.cneb_u_toit,       "cneb_u_toit",       "W/m².K"),
            criteria_row("U fenêtre — Performance des vitrages",document.cneb_u_fenetre,    "cneb_u_fenetre",    "W/m².K"),
            criteria_row("Infiltration — Étanchéité à l'air",   document.cneb_infiltration, "cneb_infiltration", "L/s.m²"),
        ])
    elif norme == 'LENOZ':
        story += criteria_section("LENOZ", [
            criteria_row("Énergie primaire",               document.lenoz_ep,     "lenoz_ep",     "kWh/m².an"),
            criteria_row("Ew — Indicateur de performance", document.lenoz_ew,     "lenoz_ew"),
            criteria_row("U mur",                          document.lenoz_u_mur,  "lenoz_u_mur",  "W/m².K"),
            criteria_row("U toit",                         document.lenoz_u_toit, "lenoz_u_toit", "W/m².K"),
        ])

    story.append(PageBreak())

    # Recommandations
    story += section_title("4", "Recommandations & points d'attention")

    if is_conform is True:
        story += reco_block(
            ">>", f"Dossier conforme aux exigences {norme}",
            "L'ensemble des critères analysés respecte les seuils réglementaires en vigueur. "
            "Aucune action corrective n'est requise pour l'obtention de la conformité.",
            GREEN_L, GREEN,
        )
    elif is_conform is None:
        story += reco_block(
            "--", "Analyse en cours",
            "Les données nécessaires à l'évaluation complète n'ont pas encore été renseignées. "
            "Les recommandations seront disponibles une fois l'analyse finalisée.",
            LGRAY, MUTED,
        )

    if norme == 'RT2012':
        if document.rt2012_bbio and document.rt2012_bbio > seuils.get('rt2012_bbio', 9999):
            story += reco_block("[!]", "Bbio — Besoins bioclimatiques non conformes",
                                "Améliorer l'isolation de l'enveloppe, optimiser l'orientation et les "
                                "surfaces vitrées, renforcer la compacité du bâtiment.", RED_L, RED)
        if document.rt2012_cep and document.rt2012_cep > seuils.get('rt2012_cep', 9999):
            story += reco_block("[!]", "Cep — Consommation énergétique non conforme",
                                "Optimiser les systèmes de chauffage, installer des équipements haute "
                                "efficacité, intégrer des énergies renouvelables.", RED_L, RED)
        if document.rt2012_tic and document.rt2012_tic > seuils.get('rt2012_tic', 9999):
            story += reco_block("[~]", "Tic — Température intérieure conventionnelle élevée",
                                "Renforcer la protection solaire, améliorer l'inertie thermique, "
                                "prévoir une ventilation nocturne efficace.",
                                colors.HexColor('#FFFBF0'), GOLD)
        if document.rt2012_airtightness and document.rt2012_airtightness > seuils.get('rt2012_airtightness', 9999):
            story += reco_block("[!]", "Étanchéité à l'air insuffisante",
                                "Revoir les jonctions et points singuliers de l'enveloppe, "
                                "traiter les passages de réseaux, réaliser un test d'infiltrométrie.", RED_L, RED)
    elif norme == 'RE2020':
        if document.re2020_energy_efficiency and document.re2020_energy_efficiency > seuils.get('re2020_energy_efficiency', 9999):
            story += reco_block("[!]", "Cep,nr — Énergie non renouvelable excessive",
                                "Privilégier des énergies décarbonées (PAC, solaire thermique), améliorer "
                                "l'isolation et réduire les consommations auxiliaires.", RED_L, RED)
        if document.re2020_carbon_emissions and document.re2020_carbon_emissions > seuils.get('re2020_carbon_emissions', 9999):
            story += reco_block("[!]", "Ic énergie — Émissions carbone non conformes",
                                "Basculer vers des énergies renouvelables, remplacer les systèmes à "
                                "combustibles fossiles, optimiser la consommation globale.", RED_L, RED)
        if document.re2020_thermal_comfort and document.re2020_thermal_comfort > seuils.get('re2020_thermal_comfort', 9999):
            story += reco_block("[~]", "DH — Confort d'été insuffisant",
                                "Installer des brise-soleils, augmenter l'inertie thermique, "
                                "prévoir une ventilation nocturne.",
                                colors.HexColor('#FFFBF0'), GOLD)
    elif norme == 'PEB':
        if document.peb_espec and document.peb_espec > seuils.get('peb_espec', 9999):
            story += reco_block("[!]", "Espec — Énergie spécifique non conforme (PEB)",
                                "Améliorer l'isolation globale, optimiser les systèmes de chauffage "
                                "et ventilation, recourir aux énergies renouvelables.", RED_L, RED)
        if document.peb_u_mur and document.peb_u_mur > seuils.get('peb_u_mur', 9999):
            story += reco_block("[!]", "U mur — Isolation des parois insuffisante",
                                "Renforcer l'isolation des murs par l'intérieur ou l'extérieur "
                                "pour atteindre le coefficient U requis par la réglementation PEB.", RED_L, RED)
    elif norme == 'MINERGIE':
        if document.minergie_qh and document.minergie_qh > seuils.get('minergie_qh', 9999):
            story += reco_block("[!]", "Qh — Besoins de chaleur trop élevés (Minergie)",
                                "Améliorer l'isolation de l'enveloppe, optimiser les vitrages "
                                "et réduire les ponts thermiques.", RED_L, RED)
        if document.minergie_n50 and document.minergie_n50 > seuils.get('minergie_n50', 9999):
            story += reco_block("[!]", "n50 — Étanchéité à l'air insuffisante (Minergie)",
                                "Traiter les points singuliers, mettre en place une membrane d'étanchéité continue.",
                                RED_L, RED)
    elif norme in ('CNEB2015', 'CNEB2020'):
        if document.cneb_ei and document.cneb_ei > seuils.get('cneb_ei', 9999):
            story += reco_block("[!]", f"Intensité énergétique non conforme ({norme})",
                                "Réduire les besoins en chauffage et climatisation, améliorer l'enveloppe "
                                "thermique, intégrer des systèmes à haute efficacité.", RED_L, RED)
    elif norme == 'LENOZ':
        if document.lenoz_ep and document.lenoz_ep > seuils.get('lenoz_ep', 9999):
            story += reco_block("[!]", "Énergie primaire non conforme (LENOZ)",
                                "Optimiser les systèmes énergétiques, intégrer des sources renouvelables "
                                "et améliorer l'enveloppe thermique.", RED_L, RED)

    # Notes admin
    if document.admin_notes:
        story += section_title("5", "Notes & observations de l'expert")
        notes_t = Table([[
            Paragraph(
                document.admin_notes.replace('\n', '<br/>'),
                st('nt', fontSize=9, leading=14),
            )
        ]], colWidths=[W])
        notes_t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), LGRAY),
            ('LINEBEFORE',    (0, 0), (0, -1), 3, GOLD),
            ('TOPPADDING',    (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING',   (0, 0), (-1, -1), 12),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 12),
        ]))
        story.append(notes_t)

    story.append(PageBreak())

    # Mentions légales
    story += section_title("6", "Mentions légales & disclaimer")
    disc_items = [
        ("Nature du rapport",
         "Ce rapport est établi sur la base des documents fournis par le client et constitue une analyse "
         "documentaire indépendante. Il ne se substitue pas à une attestation officielle de conformité "
         "délivrée par un organisme accrédité."),
        ("Responsabilité",
         "ConformExpert s'engage à fournir une analyse rigoureuse et objective des documents transmis. "
         "La conformité finale du bâtiment relève de la responsabilité du maître d'ouvrage et des "
         "professionnels en charge de la construction."),
        ("Confidentialité",
         "Ce document est strictement confidentiel et destiné exclusivement au client mentionné en page "
         "de couverture. Toute reproduction ou diffusion sans autorisation écrite de ConformExpert est interdite."),
        ("Réglementations",
         "RT2012 : Arrêté du 26 octobre 2010  |  RE2020 : Décret n°2021-1004 du 29 juillet 2021  |  "
         "PEB : Directive européenne 2010/31/UE  |  Minergie / SIA380 : Normes SIA Suisse  |  "
         "CNEB : Code national de l'énergie pour les bâtiments (Canada)  |  "
         "LENOZ : Règlement grand-ducal du 23 juillet 2016 (Luxembourg)"),
        ("Contact",
         "ConformExpert  ·  contact@conformexpert.fr  ·  Délai garanti 10 jours ouvrés"),
    ]
    for k, v in disc_items:
        disc_t = Table([[
            Paragraph(k, st('dk', fontName='Helvetica-Bold', fontSize=9, textColor=TEXT)),
            Paragraph(v, st('dv', fontSize=8.5, textColor=colors.HexColor('#444455'), leading=13)),
        ]], colWidths=[4 * cm, W - 4 * cm])
        disc_t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), LGRAY),
            ('TOPPADDING',    (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING',   (0, 0), (-1, -1), 10),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 10),
            ('LINEBELOW',     (0, 0), (-1, -1), 0.5, MGRAY),
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(disc_t)

    # Build
    buffer = BytesIO()
    doc_pdf = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=ML, rightMargin=MR,
        topMargin=MT, bottomMargin=MB,
        title=f"Rapport ConformExpert - {document.name}",
    )
    doc_pdf.build(story, onFirstPage=draw_cover, onLaterPages=draw_page)

    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    safe_name = document.name.replace(' ', '_').replace('/', '_')
    response['Content-Disposition'] = f'inline; filename="rapport_{safe_name}.pdf"'
    return response


# ──────────────────────────────────────────────────────────────
# RAPPORT IA (Claude)
# ──────────────────────────────────────────────────────────────

# Seuils réglementaires résumés par norme
_SEUILS_LABELS = {
    'RT2012':  "Bbio ≤ 60 | Cep ≤ 50 kWh ep/m².an | Tic ≤ 27°C | Étanchéité ≤ 0,6 m³/h.m²",
    'RE2020':  "Cep,nr ≤ 100 kWh/m².an | Ic énergie ≤ 160 kgCO₂/m².an | DH ≤ 1250 (zone H2)",
    'PEB':     "Espec ≤ 100 kWh/m².an | U mur ≤ 0,24 | U toit ≤ 0,20 | U plancher ≤ 0,30 W/m².K",
    'MINERGIE':"Qh ≤ 60 kWh/m².an | Qtot ≤ 38 kWh/m².an | n50 ≤ 0,6 h⁻¹",
    'SIA380':  "Qh ≤ 90 kWh/m².an selon SIA 380/1",
    'CNEB2015':"EI ≤ 170 kWh/m².an | U mur ≤ 0,24 | U toit ≤ 0,18 | U fenêtre ≤ 1,8 W/m².K",
    'CNEB2020':"EI ≤ 150 kWh/m².an | U mur ≤ 0,21 | U toit ≤ 0,16 | U fenêtre ≤ 1,6 W/m².K",
    'LENOZ':   "EP ≤ 90 kWh/m².an | Ew ≤ 100 | U mur ≤ 0,22 | U toit ≤ 0,17 W/m².K",
}

_CHAMPS_NORME = {
    'RT2012': [
        ('rt2012_bbio', 'Bbio', ''), ('rt2012_cep', 'Cep', 'kWh ep/m².an'),
        ('rt2012_tic', 'Tic', '°C'), ('rt2012_airtightness', 'Étanchéité', 'm³/h.m²'),
        ('rt2012_enr', 'ENR', ''),
    ],
    'RE2020': [
        ('re2020_energy_efficiency', 'Cep,nr', 'kWh/m².an'),
        ('re2020_carbon_emissions', 'Ic énergie CO₂', 'kgCO₂eq/m².an'),
        ('re2020_thermal_comfort', 'DH (confort été)', 'DH'),
    ],
    'PEB': [
        ('peb_espec', 'Espec', 'kWh/m².an'), ('peb_ew', 'Ew', ''),
        ('peb_u_mur', 'U mur', 'W/m².K'), ('peb_u_toit', 'U toit', 'W/m².K'),
        ('peb_u_plancher', 'U plancher', 'W/m².K'),
    ],
    'MINERGIE': [
        ('minergie_qh', 'Qh', 'kWh/m².an'), ('minergie_qtot', 'Qtot', 'kWh/m².an'),
        ('minergie_n50', 'n50', 'h⁻¹'),
    ],
    'SIA380': [('sia380_qh', 'Qh', 'kWh/m².an')],
    'CNEB2015': [
        ('cneb_ei', 'Intensité énergétique', 'kWh/m².an'),
        ('cneb_u_mur', 'U mur', 'W/m².K'), ('cneb_u_toit', 'U toit', 'W/m².K'),
        ('cneb_u_fenetre', 'U fenêtre', 'W/m².K'), ('cneb_infiltration', 'Infiltration', 'L/s.m²'),
    ],
    'CNEB2020': [
        ('cneb_ei', 'Intensité énergétique', 'kWh/m².an'),
        ('cneb_u_mur', 'U mur', 'W/m².K'), ('cneb_u_toit', 'U toit', 'W/m².K'),
        ('cneb_u_fenetre', 'U fenêtre', 'W/m².K'), ('cneb_infiltration', 'Infiltration', 'L/s.m²'),
    ],
    'LENOZ': [
        ('lenoz_ep', 'Énergie primaire', 'kWh/m².an'), ('lenoz_ew', 'Ew', ''),
        ('lenoz_u_mur', 'U mur', 'W/m².K'), ('lenoz_u_toit', 'U toit', 'W/m².K'),
    ],
}


def _build_system_prompt(type_analyse, ref, document, infos_batiment, source_donnees,
                          valeurs_str, norme, carbone_seuils_str):
    """Construit le system prompt Claude selon le type d'analyse."""

    seuils_str = _SEUILS_LABELS.get(norme, "Voir réglementation applicable")

    if type_analyse == 'carbone':
        return f"""Tu es ConformExpert, un tiers expert indépendant spécialisé en bilan carbone immobilier.
Tu réalises des bilans carbone indépendants à partir des documents fournis (rapport thermique, DPE, factures énergie, attestation RE2020, FDES).
Ton rôle est d'évaluer les émissions de CO2 du bâtiment et d'identifier les leviers de réduction.

Contexte du dossier :
- Référence : {ref}
- Projet : {document.name}
- Client : {document.client_name or 'Non renseigné'}
- Type de mission : Bilan carbone immobilier indépendant
- Informations du bâtiment :
{infos_batiment}
- Source des documents : {source_donnees}
- Référentiels applicables :
{carbone_seuils_str}

Ta mission d'expert indépendant :
1. Évaluer les émissions de CO2 liées à l'énergie du bâtiment (chauffage, électricité, eau chaude sanitaire)
2. Analyser les indicateurs Ic énergie et Ic construction si disponibles (RE2020)
3. Vérifier la conformité au regard du décret tertiaire si applicable
4. Identifier les postes d'émission prioritaires et les leviers de réduction
5. Donner un avis indépendant sur la performance carbone du bâtiment

Réponds UNIQUEMENT en JSON valide, sans markdown, sans explication, sans balises.

Structure JSON attendue :
{{
  "verdict": "Faible empreinte carbone" | "Empreinte carbone modérée" | "Empreinte carbone élevée" | "Données insuffisantes",
  "score_global": 72,
  "resume_executif": "Paragraphe de 3-5 phrases résumant la performance carbone du bâtiment, les principaux postes d'émission et les enjeux.",
  "emissions": {{
    "ic_energie": {{"valeur": null, "seuil_re2020": null, "conforme": null, "unite": "kgCO2eq/m².an"}},
    "ic_construction": {{"valeur": null, "seuil_re2020": null, "conforme": null, "unite": "kgCO2eq/m²"}},
    "conso_ep": {{"valeur": null, "unite": "kWh ep/m².an"}},
    "emission_ges": {{"valeur": null, "unite": "kgCO2eq/m².an"}},
    "classe_dpe": null
  }},
  "postes_emission": [
    {{
      "poste": "Chauffage",
      "part_estimee": "45%",
      "niveau": "faible" | "modéré" | "élevé",
      "observation": "Description basée sur les documents fournis."
    }}
  ],
  "conformite_reglementaire": [
    {{
      "referentiel": "RE2020 Ic énergie",
      "statut": "conforme" | "non_conforme" | "non_applicable" | "non_vérifié",
      "detail": "Explication."
    }}
  ],
  "leviers_reduction": [
    {{
      "levier": "Isolation thermique",
      "impact": "faible" | "modéré" | "élevé",
      "horizon": "Court terme" | "Moyen terme" | "Long terme",
      "gain_estime": "ex: -15% émissions",
      "description": "Description de l'action recommandée."
    }}
  ],
  "risques": [
    {{
      "titre": "...",
      "description": "...",
      "gravite": "faible" | "modéré" | "élevé",
      "action": "Action recommandée.",
      "urgence": "Immédiat" | "Court terme" | "Moyen terme"
    }}
  ],
  "points_forts": ["Point fort identifié"],
  "verifications_complementaires": [
    "Audit énergétique réglementaire recommandé si copropriété > 50 lots",
    "Vérification FDES matériaux si RE2020 applicable"
  ],
  "avis_independant": "Paragraphe conclusif exprimant l'avis de ConformExpert sur la performance carbone du bâtiment.",
  "mentions_legales": "Ce rapport constitue un bilan carbone documentaire indépendant réalisé par ConformExpert. Il ne se substitue pas à un bilan carbone complet certifié ni à une attestation réglementaire officielle."
}}

Base ton analyse sur les documents fournis.
Si des éléments ne sont pas documentés, l'indiquer explicitement plutôt que d'inventer.
Sois précis, factuel et indépendant."""


    else:  # 'energie' (défaut)
        return f"""Tu es ConformExpert, un tiers expert indépendant spécialisé dans la validation de rapports thermiques réglementaires.
Ton rôle n'est PAS de produire une étude thermique — le bureau d'études l'a déjà fait.
Ton rôle est de vérifier, valider et commenter de manière indépendante le travail du bureau d'études.

Contexte du dossier :
- Référence : {ref}
- Projet : {document.name}
- Client : {document.client_name or 'Non renseigné'}
- Type de mission : Validation indépendante de rapport thermique
- Norme applicable : {norme}
- Logiciel utilisé par le bureau d'études : {getattr(document, 'logiciel_detecte', 'Non détecté')}
- Informations du bâtiment :
{infos_batiment}
- Source des documents : {source_donnees}
- Valeurs déclarées dans le rapport :
{valeurs_str}
- Seuils réglementaires officiels {norme} : {seuils_str}

Ta mission de validation indépendante :
1. Vérifier que les valeurs déclarées respectent bien les seuils réglementaires {norme}
2. Évaluer la cohérence globale du rapport (les valeurs sont-elles plausibles pour ce type de bâtiment ?)
3. Identifier les éventuelles anomalies, valeurs manquantes ou incohérences
4. Croiser avec les consommations réelles issues des factures si disponibles
5. Formuler un avis indépendant clair sur la fiabilité du rapport fourni
6. Recommander des vérifications complémentaires si nécessaire

Réponds UNIQUEMENT en JSON valide, sans markdown, sans explication, sans balises.

Structure JSON attendue :
{{
  "verdict": "Conforme" | "Non Conforme" | "Données insuffisantes",
  "fiabilite_rapport": "Élevée" | "Moyenne" | "Faible" | "Non évaluable",
  "score_global": 78,
  "resume_executif": "Paragraphe de 3-5 phrases exprimant l'avis indépendant sur le rapport du bureau d'études.",
  "criteres": [
    {{
      "nom": "Nom du critère",
      "valeur": 72.0,
      "seuil": 50.0,
      "unite": "kWh ep/m².an",
      "conforme": false,
      "ecart_pct": 44.0,
      "commentaire": "Analyse indépendante de ce critère — est-il plausible ? cohérent ?"
    }}
  ],
  "points_forts": ["Point fort validé indépendamment"],
  "anomalies": [
    {{
      "critere": "Nom du critère ou aspect concerné",
      "gravite": "bloquant" | "majeur" | "mineur",
      "description": "Description précise de l'anomalie ou incohérence détectée.",
      "recommendation": "Action recommandée pour lever le doute ou corriger."
    }}
  ],
  "non_conformites": [
    {{
      "critere": "...",
      "gravite": "bloquant" | "majeur" | "mineur",
      "description": "Description du dépassement de seuil constaté.",
      "action": "Action corrective recommandée.",
      "delai": "...",
      "cout_estime": "..."
    }}
  ],
  "recommandations": [
    {{
      "priorite": "URGENT" | "RECOMMANDÉ" | "OPTIONNEL",
      "titre": "Titre de la recommandation",
      "description": "Description détaillée.",
      "impact_reglementaire": "Impact sur la conformité réglementaire.",
      "delai": "..."
    }}
  ],
  "coherence_factures": {{
    "coherent": true,
    "commentaire": "Analyse de la cohérence entre les valeurs du rapport et les consommations réelles des factures."
  }},
  "analyse_enveloppe": {{"synthese": "...", "points_attention": ["..."]}},
  "systemes_energetiques": {{"synthese": "...", "equipements": [{{"poste": "Chauffage", "equipement": "...", "performance": "...", "evaluation": "..."}}]}},
  "impact_financier": {{
    "cout_non_conformite": "Estimation du surcoût lié aux non-conformités.",
    "economies_potentielles": "Économies annuelles estimées après mise en conformité.",
    "retour_investissement": "Délai de retour sur investissement estimé."
  }},
  "verifications_complementaires": [
    "Vérification recommandée si des doutes subsistent"
  ],
  "contexte_reglementaire": "Rappel des exigences {norme} applicables à ce projet.",
  "avis_independant": "Paragraphe conclusif exprimant clairement l'avis de ConformExpert sur la qualité et la fiabilité du rapport soumis.",
  "mentions_legales": "Ce rapport constitue une analyse documentaire indépendante réalisée par ConformExpert. Il ne se substitue pas à une attestation officielle de conformité."
}}

Si une valeur n'est pas disponible pour un critère, omets ce critère.
Sois précis, factuel et indépendant. Ton rôle est celui d'un auditeur externe, pas d'un co-auteur du rapport."""


@csrf_exempt
def generer_rapport_ia(request, doc_id):
    """
    Endpoint AJAX rapport IA.
    GET              → retourne le rapport mis en cache (si existant)
    POST             → génère via Claude, sauvegarde en BDD et retourne le JSON
    POST ?force=1    → force la régénération même si déjà sauvegardé
    """
    import os
    import urllib.request
    import urllib.error

    document = get_object_or_404(Document, id=doc_id)

    # Retourner le rapport en cache si disponible
    force = request.GET.get('force') == '1'
    if not force and document.rapport_ia_json:
        try:
            return JsonResponse({'success': True, 'rapport': json.loads(document.rapport_ia_json), 'cached': True})
        except Exception:
            pass  # JSON corrompu → régénération

    if request.method not in ('POST', 'GET'):
        return JsonResponse({'error': 'Méthode invalide'}, status=405)

    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
    if not ANTHROPIC_API_KEY:
        return JsonResponse({'error': 'Clé API Anthropic manquante (ANTHROPIC_API_KEY)'}, status=500)

    # ── 1. Lire les fichiers PDF ───────────────────────────────
    pdf_b64_list = []
    for doc_file in document.fichiers.all()[:3]:
        file_path = doc_file.fichier.path
        if not file_path.lower().endswith(".pdf"):
            continue
        try:
            with open(file_path, "rb") as f:
                pdf_bytes = f.read()
            if not pdf_bytes.startswith(b"%PDF"):
                print("Fichier invalide (pas un vrai PDF):", file_path)
                continue
            pdf_b64_list.append(base64.b64encode(pdf_bytes).decode("utf-8"))
        except Exception as e:
            print("Erreur lecture PDF:", e)

    # Fallback sur l'ancien champ upload
    if not pdf_b64_list and document.upload and document.upload.name:
        try:
            with open(document.upload.path, "rb") as f:
                pdf_b64_list.append(base64.b64encode(f.read()).decode("utf-8"))
        except Exception as e:
            print(f"PDF upload indisponible : {e}")

    # ── 2. Contexte commun ─────────────────────────────────────
    type_analyse = getattr(document, 'type_analyse', 'energie') or 'energie'
    norme        = document.norme
    pays_map     = {'FR': 'France', 'BE': 'Belgique', 'CH': 'Suisse', 'CA': 'Canada', 'LU': 'Luxembourg'}
    pays_label   = pays_map.get(document.pays, document.pays)
    zone         = document.climate_zone or 'H2'
    ref          = f"DOC-{document.id:04d}"
    surface      = getattr(document, 'surface_totale', None)
    annee        = getattr(document, 'annee_construction', None)
    logements    = getattr(document, 'nombre_logements', None)

    carbone_seuils_str = (
    "\n    Référentiels bilan carbone :\n"
    "    - RE2020 Ic énergie seuil max maison : 4 kgCO2eq/m².an\n"
    "    - RE2020 Ic énergie seuil max collectif : 6.5 kgCO2eq/m².an\n"
    "    - RE2020 Ic construction seuil mi-vie 2025 : 640 kgCO2eq/m²\n"
    "    - RE2020 Ic construction seuil mi-vie 2028 : 590 kgCO2eq/m²\n"
    "    - Décret tertiaire : réduction 40% en 2030, 50% en 2040, 60% en 2050\n"
    "    - DPE passoire classe F : 70 kgCO2eq/m².an\n"
    "    - DPE passoire classe G : 100 kgCO2eq/m².an\n"
)

    infos_batiment = (
        f"- Type : {document.get_building_type_display()}\n"
        f"- Pays / Zone : {pays_label} — Zone climatique {zone}\n"
        + (f"- Surface totale : {surface} m²\n" if surface else "")
        + (f"- Année de construction : {annee}\n" if annee else "")
        + (f"- Nombre de logements : {logements}\n" if logements else "")
    )

    # Valeurs thermiques connues
    valeurs_connues = {}
    for field, label, unit in _CHAMPS_NORME.get(norme, []):
        val = getattr(document, field, None)
        if val is not None:
            valeurs_connues[label] = f"{val} {unit}".strip()
    valeurs_str = '\n'.join(f"  - {k} : {v}" for k, v in valeurs_connues.items()) \
                  or "  (aucune valeur encore saisie)"

    source_donnees = (
        f"{len(pdf_b64_list)} fichier(s) PDF joint(s) + les valeurs extraites ci-dessous"
        if pdf_b64_list else
        "les valeurs extraites ci-dessous (PDF non disponible sur le serveur)"
    )

    # ── 3. Observations expert & factures ─────────────────────
    observations_expert = document.admin_notes or "Aucune observation expert fournie."

    factures_data = []
    try:
        for f in document.factures.all():
            d = f.analyse_json or {}
            factures_data.append({
                "energie":        f.type_energie,
                "periode_debut":  d.get("periode_debut"),
                "periode_fin":    d.get("periode_fin"),
                "consommation":   d.get("consommation"),
                "montant_ttc":    d.get("montant_ttc"),
                "analyse_ok":     f.analyse_ok,
            })
    except Exception as e:
        print("Erreur lecture factures:", e)

    factures_str = json.dumps(factures_data, ensure_ascii=False)

    # ── 3b. Données d'extraction automatique (parser IA) ──────
    extraction_json   = getattr(document, 'extraction_json', None) or {}
    extraction_alertes = getattr(document, 'extraction_alertes', None) or []
    logiciel_detecte  = getattr(document, 'logiciel_detecte', '') or ''
    version_norme_det = getattr(document, 'version_norme_detectee', '') or ''

    # Valeurs issues de l'extraction automatique (complement aux valeurs saisies manuellement)
    valeurs_extraites = {}
    if isinstance(extraction_json, dict):
        valeurs_brutes = extraction_json.get('valeurs', {})
        if isinstance(valeurs_brutes, dict):
            valeurs_extraites = valeurs_brutes

    # Construire la chaîne de valeurs extraites (uniquement celles absentes des valeurs manuelles)
    lignes_extraction = []
    for k, v in valeurs_extraites.items():
        label_norm = k.lower().replace('_', ' ')
        # Inclure même si déjà dans valeurs_connues — les deux sources sont utiles pour la vérification
        lignes_extraction.append(f"  - {k} : {v}")
    extraction_valeurs_str = '\n'.join(lignes_extraction) or "  (aucune valeur extraite automatiquement)"

    # Alertes d'extraction
    if isinstance(extraction_alertes, list) and extraction_alertes:
        alertes_str = '\n'.join(f"  ⚠ {a}" for a in extraction_alertes)
    else:
        alertes_str = "  (aucune alerte d'extraction)"

    # Source enrichie
    source_donnees_enrichie = source_donnees
    if logiciel_detecte:
        source_donnees_enrichie += f" — Logiciel détecté : {logiciel_detecte}"
    if version_norme_det:
        source_donnees_enrichie += f" — Norme version : {version_norme_det}"

    # ── 4. System prompt ──────────────────────────────────────
    system_prompt = _build_system_prompt(
        type_analyse, ref, document, infos_batiment,
        source_donnees_enrichie, valeurs_str, norme, carbone_seuils_str,
    )

    # ── 6. Message utilisateur ────────────────────────────────
    user_content = []
    headers_extra = {}

    if pdf_b64_list:
        for b64 in pdf_b64_list:
            user_content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
            })
        headers_extra = {"anthropic-beta": "pdfs-2024-09-25"}
        user_content.append({"type": "text", "text": f"""
Tu es ConformExpert, un tiers expert indépendant mandaté pour valider ce rapport thermique.

Lis attentivement le(s) document(s) PDF joint(s) — il s'agit du rapport thermique produit par le bureau d'études.
Ton rôle est de le valider de manière indépendante, pas de le reproduire.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INFORMATIONS DU DOSSIER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Bâtiment :
{infos_batiment}
Logiciel bureau d'études : {logiciel_detecte or 'Non détecté'}
Version norme détectée : {version_norme_det or 'Non détectée'}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
VALEURS SAISIES MANUELLEMENT (admin)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{valeurs_str}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
VALEURS EXTRAITES AUTOMATIQUEMENT (parser IA sur le PDF)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{extraction_valeurs_str}

ALERTES DÉTECTÉES LORS DE L'EXTRACTION :
{alertes_str}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OBSERVATIONS DE L'EXPERT (admin)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{observations_expert}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FACTURES ÉNERGÉTIQUES RÉELLES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{factures_str if factures_data else "Aucune facture déposée pour ce dossier."}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TA MISSION DE VALIDATION INDÉPENDANTE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Lire et analyser le rapport thermique joint (PDF)
2. Croiser les valeurs du PDF avec les valeurs extraites automatiquement — signaler tout écart
3. Vérifier que les valeurs respectent les seuils réglementaires applicables
4. Évaluer la cohérence des valeurs avec le type et l'âge du bâtiment
5. Traiter les alertes d'extraction comme des points de vigilance prioritaires
6. Croiser les consommations théoriques avec les consommations réelles des factures si disponibles
7. Intégrer les observations de l'expert admin dans l'analyse
8. Formuler un avis indépendant clair sur la qualité et la fiabilité du rapport

RÈGLES ABSOLUES
- Ne jamais inventer de données ou de valeurs
- Si une information est absente du dossier, le mentionner explicitement
- Les alertes d'extraction sont des signaux importants — les traiter sérieusement
- Adopter le point de vue d'un auditeur externe, pas d'un co-auteur

Génère le rapport de validation complet en respectant le format JSON défini dans les instructions système.
"""})
    else:
        # Fallback sans PDF — injecter TOUTES les données disponibles
        user_content.append({"type": "text", "text": f"""
Le PDF original n'est pas disponible sur le serveur.
Génère le rapport de validation complet pour le dossier {ref} (type : {type_analyse})
en te basant EXCLUSIVEMENT sur les données ci-dessous.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INFORMATIONS DU DOSSIER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Référence : {ref}
Norme : {norme}
Logiciel bureau d'études : {logiciel_detecte or 'Non détecté'}
Version norme détectée : {version_norme_det or 'Non détectée'}

Bâtiment :
{infos_batiment}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
VALEURS SAISIES MANUELLEMENT (admin)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{valeurs_str}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
VALEURS EXTRAITES AUTOMATIQUEMENT (parser IA)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{extraction_valeurs_str}

ALERTES DÉTECTÉES LORS DE L'EXTRACTION :
{alertes_str}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OBSERVATIONS DE L'EXPERT (admin)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{observations_expert}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FACTURES ÉNERGÉTIQUES RÉELLES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{factures_str if factures_data else "Aucune facture déposée pour ce dossier."}

Génère le rapport de validation complet en respectant le format JSON défini dans les instructions système.
N'invente aucune valeur. Si une donnée est absente, indique "Non disponible" dans le champ concerné.
"""})

    # ── 7. Appel API Claude ───────────────────────────────────
    try:
        payload = json.dumps({
            "model": "claude-sonnet-4-5",
            "max_tokens": 8000,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_content}],
        }).encode('utf-8')

        headers = {
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
        }
        headers.update(headers_extra)

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload, headers=headers, method="POST",
        )
        with urllib.request.urlopen(req, timeout=300) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            raw = result['content'][0]['text'].strip().replace('```json', '').replace('```', '').strip()
            rapport = json.loads(raw)

        document.rapport_ia_json = json.dumps(rapport, ensure_ascii=False)
        document.save(update_fields=['rapport_ia_json'])

        return JsonResponse({'success': True, 'rapport': rapport, 'cached': False})

    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8')
        print(f"CLAUDE API ERROR {e.code}: {body}")
        return JsonResponse({'error': f'Erreur API Claude ({e.code}) : {body[:300]}'}, status=500)
    except json.JSONDecodeError as e:
        print(f"JSON PARSE ERROR: {e}")
        return JsonResponse({'error': f'Erreur parsing JSON : {e}'}, status=500)
    except Exception as e:
        print(f"GENERER_RAPPORT_IA ERROR: {e}")
        return JsonResponse({'error': str(e)}, status=500)


# ──────────────────────────────────────────────────────────────
# RAPPORT IA — page publique client
# ──────────────────────────────────────────────────────────────

def rapport_ia_client(request, token):
    """Page publique rapport IA — accessible via lien de suivi, sans login."""
    document = get_object_or_404(Document, tracking_token=token, status='termine')
    rapport = None

    if document.rapport_ia_json:
        try:
            rapport = json.loads(document.rapport_ia_json)
        except Exception:
            rapport = None


    factures_data = []
    try:
        for f in document.factures.all():
            d = f.analyse_json or {}
            factures_data.append({
                "energie":       f.type_energie,
                "periode_debut": d.get("periode_debut"),
                "periode_fin":   d.get("periode_fin"),
                "consommation":  d.get("consommation"),
                "montant_ttc":   d.get("montant_ttc"),
                "analyse_ok":    f.analyse_ok,
            })
    except Exception as e:
        print("Erreur lecture factures:", e)

    return render(request, "main/rapport_ia_client.html", {
        "document":      document,
        "rapport":       rapport,
        "factures_data": json.dumps(factures_data, ensure_ascii=False),
        "has_factures":  len(factures_data) > 0,
    })


# ──────────────────────────────────────────────────────────────
# FACTURES ÉNERGIE
# ──────────────────────────────────────────────────────────────

_PROMPT_FACTURE = """
Tu es un expert en analyse de factures d'énergie (électricité et gaz naturel).
Analyse attentivement cette facture et extrais toutes les données disponibles.

Réponds UNIQUEMENT avec un objet JSON valide, sans texte avant ni après, sans balises markdown.
Schéma exact :

{
  "fournisseur": "nom du fournisseur (ex: Hydro-Québec, Énergir, EDF...)",
  "type_energie": "electricite" ou "gaz",
  "periode_debut": "YYYY-MM-DD",
  "periode_fin": "YYYY-MM-DD",
  "nb_jours": 30,
  "consommation": 1250.5,
  "unite": "kWh",
  "montant_ht": 145.20,
  "montant_ttc": 162.50,
  "devise": "CAD",
  "tarif": "nom du tarif (ex: Tarif D, G, DM...)",
  "puissance_souscrite_kw": null,
  "numero_client": null,
  "numero_compteur": null,
  "adresse_consommation": null,
  "cout_par_kwh": 0.115,
  "notes": "toute observation pertinente"
}

Si une valeur est absente, utilise null.
Pour cout_par_kwh : calcule-le si possible (montant_ht / consommation).
Pour le gaz : convertis en kWh équivalent si possible (1 m³ ≈ 10.55 kWh).
"""


def _analyser_facture_ia(fichier_path):
    """Envoie un PDF de facture à Claude et retourne le dict extrait."""
    import anthropic

    client = anthropic.Anthropic()
    with open(fichier_path, 'rb') as f:
        pdf_b64 = base64.standard_b64encode(f.read()).decode('utf-8')

    resp = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=1200,
        messages=[{
            "role": "user",
            "content": [
                {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
                {"type": "text", "text": _PROMPT_FACTURE},
            ],
        }],
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw)


def upload_facture(request, doc_id):
    """Upload d'une facture PDF pour un dossier."""
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Non authentifié — rechargez la page'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Méthode non autorisée'})
    try:
        document = get_object_or_404(Document, id=doc_id)
        fichier  = request.FILES.get('fichier')
        if not fichier:
            return JsonResponse({'success': False, 'error': 'Aucun fichier reçu'})
        if not fichier.name.lower().endswith('.pdf'):
            return JsonResponse({'success': False, 'error': 'Seuls les PDF sont acceptés'})
        type_energie = request.POST.get('type_energie', 'electricite')
        facture = FactureEnergie.objects.create(
            document=document,
            fichier=fichier,
            nom=fichier.name,
            type_energie=type_energie,
        )
        return JsonResponse({'success': True, 'facture_id': facture.id, 'nom': facture.nom})
    except Exception as e:
        err = str(e)
        if 'no such table' in err or 'does not exist' in err:
            return JsonResponse({
                'success': False,
                'error': 'Migration manquante — lancez : python manage.py makemigrations && python manage.py migrate',
            })
        return JsonResponse({'success': False, 'error': err})


def analyser_facture(request, facture_id):
    """Analyse IA d'une seule facture."""
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Non authentifié'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Méthode non autorisée'})
    facture = get_object_or_404(FactureEnergie, id=facture_id)
    try:
        donnees = _analyser_facture_ia(facture.fichier.path)
        facture.analyse_json  = donnees
        facture.analyse_ok    = True
        facture.analyse_error = ''
        facture.save()
        return JsonResponse({'success': True, 'donnees': donnees})
    except Exception as e:
        facture.analyse_error = str(e)
        facture.analyse_ok    = False
        facture.save()
        return JsonResponse({'success': False, 'error': str(e)})


def analyser_toutes_factures(request, doc_id):
    """Analyse IA de toutes les factures non encore analysées d'un dossier."""
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Non authentifié'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Méthode non autorisée'})
    document  = get_object_or_404(Document, id=doc_id)
    factures  = document.factures.filter(analyse_ok=False)
    resultats = []
    for facture in factures:
        try:
            donnees = _analyser_facture_ia(facture.fichier.path)
            facture.analyse_json  = donnees
            facture.analyse_ok    = True
            facture.analyse_error = ''
            facture.save()
            resultats.append({'id': facture.id, 'nom': facture.nom, 'ok': True})
        except Exception as e:
            facture.analyse_error = str(e)
            facture.analyse_ok    = False
            facture.save()
            resultats.append({'id': facture.id, 'nom': facture.nom, 'ok': False, 'error': str(e)})
    return JsonResponse({'success': True, 'resultats': resultats})


def supprimer_facture(request, facture_id):
    """Suppression d'une facture et de son fichier."""
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Non authentifié'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Méthode non autorisée'})
    facture = get_object_or_404(FactureEnergie, id=facture_id)
    try:
        facture.fichier.delete(save=False)
        facture.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


def get_donnees_factures(request, doc_id):
    """Retourne les données agrégées des factures analysées."""
    try:
        document = get_object_or_404(Document, id=doc_id)
        factures = document.factures.filter(analyse_ok=True).order_by('uploaded_at')
    except Exception as e:
        return JsonResponse({
            'success': False, 'error': f'Migration manquante : {e}',
            'mois': [], 'nb_factures': 0,
        })

    mois = []
    for f in factures:
        d = f.analyse_json or {}
        if d.get('consommation') is not None:
            mois.append({
                'periode_debut': d.get('periode_debut'),
                'periode_fin':   d.get('periode_fin'),
                'consommation':  d.get('consommation'),
                'unite':         d.get('unite', 'kWh'),
                'montant_ttc':   d.get('montant_ttc'),
                'cout_par_kwh':  d.get('cout_par_kwh'),
                'type_energie':  f.type_energie,
                'fournisseur':   d.get('fournisseur'),
                'devise':        d.get('devise', 'CAD'),
                'nom':           f.nom,
            })

    conso_totale = sum(m['consommation'] for m in mois if m['consommation'])
    cout_total   = sum(m['montant_ttc']  for m in mois if m['montant_ttc'])
    cout_moyen   = round(cout_total / conso_totale, 4) if conso_totale else None
    pic          = max((m['consommation'] for m in mois if m['consommation']), default=None)

    return JsonResponse({
        'success':        True,
        'mois':           mois,
        'conso_totale':   round(conso_totale, 1),
        'cout_total':     round(cout_total, 2),
        'cout_moyen_kwh': cout_moyen,
        'pic_mensuel':    pic,
        'nb_factures':    len(mois),
    })



# ──────────────────────────────────────────────────────────────
# ANALYSE MANUELLE D'UN DOCUMENT — endpoint AJAX
# ──────────────────────────────────────────────────────────────

@csrf_exempt
def analyser_document(request, doc_id):
    """
    Endpoint AJAX — (ré)analyse un dossier avec le parser intelligent.
    POST /dossier/<doc_id>/analyser/
    Retourne le résultat de détection + valeurs extraites + alertes.
    """
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Non authentifié'}, status=401)
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode invalide'}, status=405)

    document = get_object_or_404(Document, id=doc_id)

    # Lire tous les fichiers PDF du dossier
    texte_complet = ""
    pdf_b64_principal = None

    for doc_file in document.fichiers.all():
        try:
            texte_complet += extract_text_from_pdf(doc_file.fichier.path) + "\n\n"
            # Premier PDF valide comme document principal
            if pdf_b64_principal is None and doc_file.fichier.path.lower().endswith('.pdf'):
                with open(doc_file.fichier.path, 'rb') as f:
                    pdf_bytes = f.read()
                if pdf_bytes.startswith(b'%PDF'):
                    pdf_b64_principal = base64.b64encode(pdf_bytes).decode('utf-8')
        except Exception as e:
            print(f"Erreur lecture fichier {doc_file.nom}: {e}")

    # Fallback ancien champ upload
    if not texte_complet and document.upload and document.upload.name:
        try:
            texte_complet = extract_text_from_pdf(document.upload.path)
            with open(document.upload.path, 'rb') as f:
                pdf_bytes = f.read()
            if pdf_bytes.startswith(b'%PDF'):
                pdf_b64_principal = base64.b64encode(pdf_bytes).decode('utf-8')
        except Exception as e:
            print(f"Erreur upload principal: {e}")

    if not texte_complet and not pdf_b64_principal:
        return JsonResponse({'error': 'Aucun document PDF trouvé dans ce dossier'}, status=400)

    # Pour un bilan carbone, on ne lance pas le parser thermique
    if document.type_analyse == 'carbone':
        return JsonResponse({
            'success':              True,
            'type_rapport':         'bilan_carbone',
            'type_rapport_label':   'Bilan carbone immobilier',
            'logiciel_detecte':     None,
            'version_norme':        None,
            'norme_appliquee':      None,
            'nb_valeurs_extraites': 0,
            'alertes':              [],
            'conformite_declaree':  'non_applicable',
            'resume':               "Documents reçus pour bilan carbone. Générez le rapport IA pour lancer l'analyse.",
            'batiment':             {},
            'valeurs':              {},
        })

    try:
        resultat = analyser_rapport_thermique(texte_complet, pdf_b64=pdf_b64_principal)
        valeurs  = resultat.get('valeurs', {})
        analyze_document(document, valeurs, resultat)

        return JsonResponse({
            'success':              True,
            'type_rapport':         document.type_rapport,
            'type_rapport_label':   document.type_rapport_label,
            'logiciel_detecte':     document.logiciel_detecte,
            'version_norme':        document.version_norme_detectee,
            'norme_appliquee':      document.norme,
            'nb_valeurs_extraites': len(valeurs),
            'alertes':              document.extraction_alertes or [],
            'conformite_declaree':  resultat.get('conformite_declaree', 'non_precise'),
            'resume':               resultat.get('resume_extraction', ''),
            'batiment':             resultat.get('batiment', {}),
            'valeurs':              valeurs,
        })

    except Exception as e:
        print(f"ANALYSER_DOCUMENT ERROR: {e}")
        return JsonResponse({'error': str(e)}, status=500)

# ──────────────────────────────────────────────────────────────
# API REST
# ──────────────────────────────────────────────────────────────

@csrf_exempt
@api_view(['GET', 'POST'])
def api_document_list(request):
    if request.method == 'GET':
        documents  = Document.objects.all()
        serializer = DocumentSerializer(documents, many=True)
        return Response(serializer.data)
    serializer = DocumentSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@csrf_exempt
@api_view(['GET'])
def api_document_detail(request, pk):
    document   = get_object_or_404(Document, pk=pk)
    serializer = DocumentSerializer(document)
    return Response(serializer.data)


@csrf_exempt
@api_view(['GET'])
def api_results(request):
    documents  = Document.objects.all()
    serializer = DocumentSerializer(documents, many=True)
    return Response(serializer.data)


@csrf_exempt
@api_view(['GET'])
def api_history(request):
    documents  = Document.objects.all().order_by('-upload_date')
    serializer = DocumentSerializer(documents, many=True)
    return Response(serializer.data)


@csrf_exempt
@api_view(['GET'])
def api_report(request, pk):
    document = get_object_or_404(Document, pk=pk)
    context  = {
        'document':     document,
        're2020_limits': fetch_re2020_requirements(),
        'rt2012_limits': fetch_rt2012_requirements(),
    }
    try:
        from weasyprint import HTML as WeasyprintHTML
        html_string = render_to_string('main/report_template.html', context)
        pdf = WeasyprintHTML(string=html_string).write_pdf()
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="report_{document.name}.pdf"'
        return response
    except ImportError:
        return Response({'error': 'WeasyPrint non installé.'}, status=500)


# ──────────────────────────────────────────────────────────────
# DEVIS — gestion admin
# ──────────────────────────────────────────────────────────────

@login_required(login_url='/login/')
def devis_list(request):
    from django.db.models import Sum

    current_statut = request.GET.get('statut', '')
    qs = Devis.objects.all()
    if current_statut:
        qs = qs.filter(statut=current_statut)

    today = date.today()
    total        = Devis.objects.count()
    nb_acceptes  = Devis.objects.filter(statut='accepte').count()
    nb_attente   = Devis.objects.filter(statut='en_attente').count()
    nb_refuses   = Devis.objects.filter(statut='refuse').count()

    ca_mois    = Devis.objects.filter(statut='accepte', created_at__year=today.year, created_at__month=today.month).aggregate(s=Sum('montant'))['s'] or 0
    ca_total   = Devis.objects.filter(statut='accepte').aggregate(s=Sum('montant'))['s'] or 0
    ca_attente = Devis.objects.filter(statut='en_attente').aggregate(s=Sum('montant'))['s'] or 0

    taux_conversion = round(nb_acceptes / total * 100) if total else 0
    taux_attente    = round(nb_attente  / total * 100) if total else 0
    taux_refuses    = round(nb_refuses  / total * 100) if total else 0

    # Revenus sur 6 mois glissants
    revenus_mois = []
    max_montant  = 1
    for i in range(5, -1, -1):
        m = (today.month - i - 1) % 12 + 1
        y = today.year - ((today.month - i - 1) // 12)
        montant = Devis.objects.filter(
            statut='accepte', created_at__year=y, created_at__month=m,
        ).aggregate(s=Sum('montant'))['s'] or 0
        revenus_mois.append({'label': calendar.month_abbr[m][:3], 'montant': int(montant), 'raw': montant})
        if montant > max_montant:
            max_montant = montant
    for r in revenus_mois:
        r['pct'] = round(r['raw'] / max_montant * 100) if max_montant else 0

    return render(request, 'main/devis_list.html', {
        'devis_list':       qs,
        'total':            total,
        'nb_acceptes':      nb_acceptes,
        'nb_attente':       nb_attente,
        'nb_refuses':       nb_refuses,
        'ca_mois':          int(ca_mois),
        'ca_total':         int(ca_total),
        'ca_attente':       int(ca_attente),
        'taux_conversion':  taux_conversion,
        'taux_attente':     taux_attente,
        'taux_refuses':     taux_refuses,
        'revenus_mois':     revenus_mois,
        'current_statut':   current_statut,
        'statut_choices':   Devis.STATUT_CHOICES,
    })


def _devis_from_post(d, post):
    """Hydrate un objet Devis depuis les données POST."""
    d.client_nom    = post.get('client_nom', '').strip()
    d.client_email  = post.get('client_email', '').strip()
    d.client_phone  = post.get('client_phone', '').strip()
    d.projet_nom    = post.get('projet_nom', '').strip()
    d.type_batiment = post.get('type_batiment', 'maison')
    d.norme         = post.get('norme', 'RE2020')
    d.statut        = post.get('statut', 'en_attente')
    d.notes         = post.get('notes', '').strip()
    montant         = post.get('montant', '').strip()
    d.montant       = float(montant) if montant else None
    return d


@login_required(login_url='/login/')
def devis_create(request):
    if request.method == 'POST':
        d = _devis_from_post(Devis(), request.POST)
        d.save()
        messages.success(request, f'Devis pour {d.client_nom} créé.')
        return redirect('devis_edit', d.id)
    return render(request, 'main/devis_form.html', {'devis': None})


@login_required(login_url='/login/')
def devis_edit(request, devis_id):
    d = get_object_or_404(Devis, id=devis_id)
    if request.method == 'POST':
        _devis_from_post(d, request.POST)
        d.save()
        messages.success(request, 'Devis mis à jour.')
        return redirect('devis_edit', d.id)
    return render(request, 'main/devis_form.html', {'devis': d})


@login_required(login_url='/login/')
def devis_delete(request, devis_id):
    d = get_object_or_404(Devis, id=devis_id)
    if request.method == 'POST':
        d.delete()
        messages.success(request, 'Devis supprimé.')
    return redirect('devis_list')


# ── MESSAGERIE ────────────────────────────────────────────────────────────────

@login_required(login_url='/login/')
def admin_send_message(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    if request.method == 'POST':
        contenu = request.POST.get('contenu', '').strip()
        fichier = request.FILES.get('fichier')
        if contenu or fichier:
            msg = Message(document=document, auteur='admin', contenu=contenu)
            if fichier:
                msg.fichier = fichier
                msg.fichier_nom = fichier.name
            msg.save()
    return redirect('edit_document', doc_id=doc_id)


def client_send_message(request, token):
    document = get_object_or_404(Document, tracking_token=token)
    if request.method == 'POST':
        contenu = request.POST.get('contenu', '').strip()
        fichier = request.FILES.get('fichier')
        if contenu or fichier:
            msg = Message(document=document, auteur='client', contenu=contenu)
            if fichier:
                msg.fichier = fichier
                msg.fichier_nom = fichier.name
            msg.save()
    return redirect('tracking', token=token)
